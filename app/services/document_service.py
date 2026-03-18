from fastapi import UploadFile
from datetime import datetime
from typing import List, Optional, Dict, Any, Union
import asyncio
import json
from io import BytesIO
from collections import Counter, defaultdict
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from PIL import Image
import copy
from sqlalchemy import UUID, DateTime, and_, or_, cast, String, select, text, func
from sqlalchemy.orm import Session, joinedload

from app.models.document_share import DocumentShare
from app.models.document_type import DocumentType
from app.models.form import FormField
from app.models.organisation import Organisation

from ..models.document import Document
from ..models.status import Status
from ..models.user import User
from ..models.user_role import UserRole
from ..models.role import Role
from ..models.user_client import UserClient
from ..models.client import Client
from ..models.document_form_data import DocumentFormData
from ..services.s3_service import s3_service
from ..services.websocket_manager import websocket_manager
from ..services.webhook_service import webhook_service
from ..core.database import SessionLocal
from app.models import client
from app.models import user


class DocumentService:

    @staticmethod
    def get_status_id_by_code(db: Session, code: str) -> int:
        status = db.query(Status).filter(Status.code == code).first()
        if status:
            return status.id
        return None

    @staticmethod
    def _get_user_role_flags(user: User):
        role_names = [r.name for r in user.roles]
        return {
            "is_super_admin": "SUPER_ADMIN" in role_names,
            "is_admin": "ADMIN" in role_names,
            "is_client": user.is_client,
            "is_staff": not user.is_client
        }

    @staticmethod
    def _document_access_query(db: Session, actor: User):
        base = db.query(Document)
        if getattr(actor, "context_temp", False):
            return base.filter(text("1=0"))

        org_id = getattr(actor, "context_organisation_id", None) or getattr(actor, "organisation_id", None)
        if not org_id and hasattr(actor, "id") and not hasattr(actor, "organisation_id"):
            org_id = actor.id

        if not org_id:
            return base.filter(text("1=0"))

        role_names = [r.name for r in getattr(actor, "roles", [])]

        if getattr(actor, "is_superuser", False) or "ORGANISATION_ADMIN" in role_names:
            return base.filter(Document.organisation_id == org_id)

        if actor.is_client:
            client_docs = base.filter(
                and_(
                    Document.organisation_id == org_id,
                    or_(
                        Document.created_by == actor.id,
                        Document.client_id == actor.client_id
                    )
                )
            )
            shared_docs = (
                db.query(Document)
                .join(DocumentShare, DocumentShare.document_id == Document.id)
                .filter(
                    DocumentShare.user_id == actor.id,
                    Document.organisation_id == org_id
                )
            )
            return client_docs.union(shared_docs)

        assigned_clients = (
            db.query(UserClient.client_id)
            .filter(UserClient.user_id == actor.id)
        )

        staff_docs = base.filter(
            and_(
                Document.organisation_id == org_id,
                or_(
                    Document.created_by == actor.id,
                    Document.client_id.in_(assigned_clients)
                )
            )
        )
        shared_docs = (
            db.query(Document)
            .join(DocumentShare, DocumentShare.document_id == Document.id)
            .filter(
                DocumentShare.user_id == actor.id,
                Document.organisation_id == org_id
            )
        )
        return staff_docs.union(shared_docs)

    @staticmethod
    def _get_accessible_document(db: Session, document_id: int, user: User):
        return (
            DocumentService._document_access_query(db, user)
            .filter(Document.id == document_id)
            .first()
        )

    @staticmethod
    def get_total_pages(file_bytes: bytes, content_type: str) -> int:
        try:
            if content_type == "application/pdf":
                fp = BytesIO(file_bytes)
                parser = PDFParser(fp)
                doc = PDFDocument(parser)
                return sum(1 for _ in PDFPage.create_pages(doc))

            if content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                from docx import Document as DocxDocument
                doc = DocxDocument(BytesIO(file_bytes))
                return max(1, len(doc.sections))

            if content_type in ("image/png", "image/jpeg", "image/jpg"):
                return 1

            if content_type == "image/tiff":
                img = Image.open(BytesIO(file_bytes))
                return getattr(img, "n_frames", 1)

        except Exception:
            pass
        return 1

    @staticmethod
    def build_derived_document_counts(extracted_docs, unverified_docs):
        counts = defaultdict(int)
        for ed in extracted_docs:
            if ed.document_type:
                counts[ed.document_type.name] += 1
        for ud in unverified_docs:
            if ud.suspected_type:
                counts[ud.suspected_type] += 1
        return dict(counts)

    # ─────────────────────────────────────────────────────────────────────────
    # FIX [5]: extract_client_id_from_form_data — method was missing entirely
    # Called by the router's PATCH /form-data endpoint → AttributeError crash
    # ─────────────────────────────────────────────────────────────────────────
    @staticmethod
    def extract_client_id_from_form_data(db: Session, form_data: dict):
        """
        Extract client_id from form_data dict.
        Checks both the direct 'client_id' key and any FormField whose label == 'Client'.
        Returns the client UUID or None.
        """
        if not form_data:
            return None

        # Direct key (set by enforcement logic)
        if "client_id" in form_data:
            return form_data["client_id"]

        # Fallback: scan FormField labels
        for field_id, value in form_data.items():
            if not value:
                continue
            field = db.query(FormField).filter(FormField.id == field_id).first()
            if field and field.label.lower() == "client":
                return value

        return None

    @staticmethod
    async def create_document(db: Session, file: UploadFile, user) -> Document:
        status_id = DocumentService.get_status_id_by_code(db, "QUEUED")
        if isinstance(user, User):
            created_by = user.id
            org_id = getattr(user, "context_organisation_id", None)
        document = Document(
            filename=file.filename,
            original_filename=file.filename,
            file_size=file.size,
            content_type=file.content_type,
            created_by=user.id,
            organisation_id=org_id,
            status_id=status_id
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        return document

    @staticmethod
    async def update_document_status(db: Session, document_id: int, status_code: str,
                                     progress: int = 0, error_message: str = None,
                                     s3_key: str = None, s3_bucket: str = None):
        document = db.query(Document).filter(Document.id == document_id).first()
        if document:
            status_id = DocumentService.get_status_id_by_code(db, status_code)
            if status_id:
                document.status_id = status_id
            document.upload_progress = progress
            if error_message:
                document.error_message = error_message
            if s3_key:
                document.s3_key = s3_key
            if s3_bucket:
                document.s3_bucket = s3_bucket

            db.commit()

            await websocket_manager.broadcast_document_status(
                document_id=document.id,
                status=status_code,
                user_id=str(document.created_by),
                progress=progress,
                error_message=error_message
            )

    @staticmethod
    async def upload_document_to_s3(db: Session, document_id: int, file: UploadFile):
        try:
            await DocumentService.update_document_status(db, document_id, "UPLOADING", progress=0)
            s3_key, bucket_name = await s3_service.upload_file(
                file.file, file.filename, file.content_type
            )
            await DocumentService.update_document_status(
                db, document_id, "UPLOADED",
                progress=100, s3_key=s3_key, s3_bucket=bucket_name
            )
        except Exception as e:
            await DocumentService.update_document_status(
                db, document_id, "UPLOAD_FAILED", error_message=str(e)
            )
            raise e

    @staticmethod
    async def process_multiple_uploads(db: Session, files: List[UploadFile], user: User,
                                       enable_ai: bool = False, document_type_id: str = None,
                                       template_id: str = None, form_id: str = None,
                                       form_data: str = None):
        documents = []
        file_buffers = []
        from ..models.document_form_data import DocumentFormData

        queued_status_id = DocumentService.get_status_id_by_code(db, "QUEUED")

        parsed_form_data = json.loads(form_data) if form_data else {}

        if isinstance(user, User) and user.is_client:
            parsed_form_data["client_id"] = str(user.client_id)
            if form_id:
                form_fields = (
                    db.query(FormField)
                    .filter(FormField.form_id == str(form_id))
                    .all()
                )
                for field in form_fields:
                    label = field.label.lower()
                    if label == "client":
                        continue
                    if field.default_value is not None:
                        field_key = str(field.id)
                        if field_key not in parsed_form_data:
                            parsed_form_data[field_key] = field.default_value

        if isinstance(user, User) and user.is_client:
            client_id_value = user.client_id
            parsed_form_data["client_id"] = str(user.client_id)
        else:
            client_id_value = parsed_form_data.get("client_id")

        for file in files:
            content = await file.read()
            file_size = len(content)
            buffer = BytesIO(content)
            total_pages = DocumentService.get_total_pages(content, file.content_type)
            await file.seek(0)

            if isinstance(user, Organisation):
                created_by = None
                org_id = user.id
                client_id_value = parsed_form_data.get("client_id")
            elif isinstance(user, User):
                created_by = user.id
                org_id = getattr(user, "context_organisation_id", None)
                client_id_value = parsed_form_data.get("client_id") or user.client_id
            else:
                raise Exception("Unknown actor")

            document = Document(
                filename=file.filename,
                original_filename=file.filename,
                file_size=file_size,
                content_type=file.content_type,
                created_by=created_by,
                organisation_id=org_id,
                client_id=client_id_value,
                status_id=queued_status_id,
                upload_progress=0,
                enable_ai=enable_ai,
                document_type_id=document_type_id,
                template_id=template_id,
                total_pages=total_pages
            )
            db.add(document)
            db.flush()
            documents.append(document)

            if parsed_form_data or form_id:
                # BUG FIX: copy the dict so each document gets its own independent
                # data object. Without this, all documents in a batch upload share
                # the same Python dict reference — editing one mutates all others.
                form_data_record = DocumentFormData(
                    document_id=document.id,
                    form_id=form_id,
                    data=copy.deepcopy(parsed_form_data)  # deep copy — form_data may contain nested dicts
                )
                db.add(form_data_record)

            file_buffers.append({
                'buffer': buffer,
                'filename': file.filename,
                'content_type': file.content_type
            })

        db.commit()

        for document in documents:
            await websocket_manager.broadcast_document_status(
                document_id=document.id,
                status="QUEUED",
                user_id=str(document.created_by),
                progress=0
            )

        asyncio.create_task(DocumentService._process_uploads_background(
            documents, file_buffers, enable_ai, document_type_id, template_id
        ))

        return documents

    @staticmethod
    async def _process_uploads_background(documents: List[Document], files_data: List[dict],
                                          enable_ai: bool = False, document_type_id: str = None,
                                          template_id: str = None):
        try:
            upload_tasks = [
                DocumentService._process_single_upload_only(doc.id, file_data)
                for doc, file_data in zip(documents, files_data)
            ]
            results = await asyncio.gather(*upload_tasks, return_exceptions=True)

            successful_docs = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    print(f"Upload failed for document {documents[i].id}: {result}")
                else:
                    successful_docs.append((documents[i], files_data[i]))

            if enable_ai and successful_docs:
                db = SessionLocal()
                try:
                    for doc, file_data in successful_docs:
                        await DocumentService.update_document_status(
                            db, doc.id, "AI_QUEUED",
                            progress=0, error_message="Queued for AI Analysis"
                        )
                finally:
                    db.close()

                for doc, file_data in successful_docs:
                    await DocumentService._process_single_ai_analysis(
                        doc.id, file_data, document_type_id, template_id
                    )

        except Exception as e:
            print(f"Error in background processing: {e}")

    @staticmethod
    async def _process_single_upload_only(document_id: int, file_data: dict):
        db = SessionLocal()
        try:
            await DocumentService.update_document_status(db, document_id, "UPLOADING", progress=10)

            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise Exception("Document not found")

            file_bytes = file_data['buffer'].getvalue()
            total_size = len(file_bytes)
            uploaded_bytes = 0
            main_loop = asyncio.get_event_loop()
            file_data['buffer'].seek(0)

            def progress_callback(bytes_amount):
                nonlocal uploaded_bytes
                uploaded_bytes += bytes_amount
                percentage = min(int((uploaded_bytes / total_size) * 90) + 10, 99)
                if document and document.created_by:
                    asyncio.run_coroutine_threadsafe(
                        websocket_manager.broadcast_document_status(
                            document_id=document_id,
                            status="UPLOADING",
                            user_id=str(document.created_by),
                            progress=percentage
                        ),
                        main_loop
                    )

            safe_filename = "".join(
                c for c in file_data['filename'] if c.isalnum() or c in ('._-')
            ).strip() or f"doc_{document_id}"

            custom_s3_key = f"documents/{document.created_by}/{document_id}_{safe_filename}"
            s3_buffer = BytesIO(file_bytes)

            s3_key, bucket_name = await s3_service.upload_file(
                s3_buffer,
                file_data['filename'],
                file_data['content_type'],
                progress_callback=progress_callback,
                s3_key=custom_s3_key
            )

            await DocumentService.update_document_status(
                db, document_id, "UPLOADED",
                progress=100, s3_key=s3_key, s3_bucket=bucket_name
            )

            asyncio.create_task(asyncio.to_thread(
                webhook_service.trigger_webhook_background,
                "document.uploaded",
                {"document_id": document_id, "filename": document.filename, "s3_key": s3_key},
                str(document.created_by),
                SessionLocal
            ))

            return True

        except Exception as e:
            await DocumentService.update_document_status(
                db, document_id, "UPLOAD_FAILED", error_message=str(e)
            )
            raise e
        finally:
            db.close()

    @staticmethod
    async def _process_single_ai_analysis(document_id: int, file_data: dict,
                                           document_type_id: str = None,
                                           template_id: str = None,
                                           analysis_result=None):
        db = SessionLocal()
        try:
            await DocumentService.update_document_status(
                db, document_id, "ANALYZING",
                progress=0, error_message="Starting Analysis..."
            )

            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                raise Exception("Document not found")

            async def report_ai_progress(msg, pct):
                await DocumentService.update_document_status(
                    db, document_id, "ANALYZING", progress=pct, error_message=msg
                )

            file_bytes = file_data['buffer'].getvalue()

            from ..models.template import Template
            from ..models.document_type import DocumentType
            from ..models.extracted_document import ExtractedDocument
            from ..models.unverified_document import UnverifiedDocument
            from ..services.ai_service import AIService

            active_status = db.query(Status).filter(Status.code == "ACTIVE").first()
            if not active_status:
                raise Exception("ACTIVE status not found")

            # Fetch ACTIVE doc types only — inactive types must not appear
            # in the classifier prompt (they confuse the model with empty descriptions)
            doc_types = db.query(DocumentType).filter(
                DocumentType.organisation_id == document.organisation_id,
                DocumentType.status_id == active_status.id,
            ).all()

            doc_type_map = {dt.name.strip().upper(): dt for dt in doc_types}

            # ── Build schemas ──────────────────────────────────────────────────
            # ALL document types for this org are passed to ai_service so the
            # classifier sees every type the org has defined — not just those
            # that happen to have an active template.
            #
            # Two tiers:
            #   WITH active template  → classifier + field extraction → ExtractedDocument
            #   WITHOUT template      → classifier only → saved as UnverifiedDocument
            #                           (staff can review / add a template later)
            # ──────────────────────────────────────────────────────────────────
            templates = db.query(Template).filter(
                Template.status_id == active_status.id,
                Template.organisation_id == document.organisation_id
            ).all()

            template_group = defaultdict(list)
            for t in templates:
                template_group[t.document_type_id].append(t)

            schemas = []
            active_template_map = {}

            for dt in doc_types:
                normalized_name = dt.name.strip().upper()
                grouped = template_group.get(dt.id, [])
                description = (dt.description or "") if hasattr(dt, "description") else ""

                if not grouped:
                    # No active template — include in classifier but with no fields.
                    # ai_service will classify pages as this type and save as UnverifiedDocument.
                    schemas.append({
                        "type_name":   normalized_name,
                        "description": description,
                        "fields":      [],
                    })
                    continue

                # Only ONE active template per doc type enforced
                if len(grouped) > 1:
                    raise Exception(
                        f"Multiple ACTIVE templates found for document type '{dt.name}'. "
                        "Deactivate all but one before re-analysing."
                    )

                template_obj = grouped[0]
                schemas.append({
                    "type_name":   normalized_name,
                    "description": description,
                    "fields":      template_obj.extraction_fields or [],
                })
                # Only register in active_template_map if it actually has fields to extract
                if template_obj.extraction_fields:
                    active_template_map[normalized_name] = template_obj

            print(f"[AI] document_id={document_id} "
                  f"classifier_types={[s['type_name'] for s in schemas]} "
                  f"extractable={list(active_template_map.keys())}")

            async def check_cancelled():
                check_db = SessionLocal()
                try:
                    current_doc = check_db.query(Document).filter(Document.id == document_id).first()
                    cancelled_status_id = DocumentService.get_status_id_by_code(check_db, "CANCELLED")
                    return current_doc and current_doc.status_id == cancelled_status_id
                finally:
                    check_db.close()

            # Instantiate per-call so each analysis gets a fresh OpenAI client
            _ai = AIService()
            analysis_result = await _ai.analyze_document(
                file_bytes,
                file_data['filename'],
                schemas,
                db,
                document.id,
                progress_callback=report_ai_progress,
                check_cancelled_callback=check_cancelled
            )

            findings = analysis_result.get("findings", [])

            # ─────────────────────────────────────────────────────────────────
            # FIX [8]: normalize_fields — the list branch used item.get("exampleValue")
            # which is a SCHEMA field, not a runtime value.
            # ai_service returns data={"fields": {fieldName: value}} flat dict,
            # so we only ever hit the dict branch in practice. The list branch
            # is kept for safety but now uses item.get("value") as fallback.
            # ─────────────────────────────────────────────────────────────────
            def normalize_fields(raw_fields):
                if isinstance(raw_fields, dict):
                    return raw_fields
                if isinstance(raw_fields, list):
                    normalized = {}
                    for item in raw_fields:
                        if not isinstance(item, dict):
                            continue
                        field_name = item.get("fieldName")
                        # "value" / "extractedValue" are runtime keys
                        # "exampleValue" is a schema-definition key — WRONG for runtime use
                        value = item.get("value") or item.get("extractedValue")  # not exampleValue
                        if field_name:
                            normalized[field_name] = value
                    return normalized
                return {}

            # ─────────────────────────────────────────────────────────────────
            # AIService.analyze_document already:
            #   1. Hard-resets UnverifiedDocuments at the START (idempotency)
            #   2. Saves UnverifiedDocuments DURING processing (unknown types)
            #   3. Returns findings containing ONLY schema-matched (verified) docs
            #
            # So we must ONLY clear ExtractedDocuments here (verified results from
            # any prior run). Never delete UnverifiedDocuments — AIService already
            # wrote them and deleting them would silently lose that data.
            # ─────────────────────────────────────────────────────────────────
            db.query(ExtractedDocument).filter(
                ExtractedDocument.document_id == document.id
            ).delete()
            db.commit()

            excel_rows = []

            for finding in findings:
                doc_type_raw = finding.get("type", "")
                doc_type     = doc_type_raw.strip().upper()
                page_range   = finding.get("page_range")
                confidence   = finding.get("confidence", 1.0)

                if not doc_type or not page_range:
                    continue

                # findings[].data = {"fields": {fieldName: value}} from AIService
                raw_fields = finding.get("data", {}).get("fields", {})
                fields     = normalize_fields(raw_fields)

                doc_type_obj = doc_type_map.get(doc_type)
                template_obj = active_template_map.get(doc_type)

                if template_obj:
                    # Whitelist to only template-defined fields
                    expected_fields = [f["fieldName"] for f in template_obj.extraction_fields]
                    fields = {ef: fields.get(ef) for ef in expected_fields}

                    row = {"Document Type": doc_type, "Page Range": page_range}
                    row.update(fields)
                    excel_rows.append(row)

                if doc_type_obj and template_obj:
                    # Verified: matched both a DocumentType and an active Template
                    db.add(ExtractedDocument(
                        document_id=document.id,
                        document_type_id=doc_type_obj.id,
                        template_id=template_obj.id,
                        extracted_data=fields,
                        page_range=page_range,
                        confidence=confidence
                    ))
                else:
                    # AIService returned a type it classified but we have no DB template for —
                    # save as unverified so staff can review (avoids silent data loss)
                    db.add(UnverifiedDocument(
                        document_id=document.id,
                        suspected_type=doc_type_raw,
                        page_range=page_range,
                        extracted_data=fields,
                        status="PENDING"
                    ))

            db.commit()

            # Excel report
            if excel_rows:
                import pandas as pd
                import io as _io
                df = pd.DataFrame(excel_rows)
                excel_buffer = _io.BytesIO()
                with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name='Analysis Report')
                excel_buffer.seek(0)

                report_s3_key, _ = await s3_service.upload_file(
                    excel_buffer,
                    f"analysis_report_{document_id}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
                document.analysis_report_s3_key = report_s3_key
                db.commit()

            # Error check — only flag _error entries as failures
            error_findings = [
                f for f in findings
                if f.get("data", {}).get("_error")
            ]

            if error_findings:
                first_error = error_findings[0]["data"]["_error"]
                await DocumentService.update_document_status(
                    db, document_id, "AI_FAILED",
                    progress=100,
                    error_message=f"Partial Analysis Failure: {first_error}"
                )
                asyncio.create_task(asyncio.to_thread(
                    webhook_service.trigger_webhook_background,
                    "document.failed",
                    {"document_id": document_id, "filename": document.filename,
                     "error": f"Partial Analysis Failure: {first_error}"},
                    str(document.created_by),
                    SessionLocal
                ))
            else:
                await DocumentService.update_document_status(
                    db, document_id, "COMPLETED",
                    progress=100, error_message="Analysis Complete"
                )
                asyncio.create_task(asyncio.to_thread(
                    webhook_service.trigger_webhook_background,
                    "document.processed",
                    {"document_id": document_id, "filename": document.filename, "status": "COMPLETED"},
                    str(document.created_by),
                    SessionLocal
                ))

        except Exception as e:
            import traceback as _tb
            error_str = str(e)
            is_cancelled = "Analysis Cancelled" in error_str
            if not is_cancelled:
                print(f"[AI] document_id={document_id} FAILED: {error_str}")
                _tb.print_exc()

            await DocumentService.update_document_status(
                db, document_id,
                "CANCELLED" if is_cancelled else "AI_FAILED",
                error_message="Analysis Cancelled" if is_cancelled else error_str
            )

            if not is_cancelled:
                # ─────────────────────────────────────────────────────────────
                # FIX [7]: safe None-guard — don't chain .filename / .created_by
                # on a potentially-None query result
                # ─────────────────────────────────────────────────────────────
                failed_doc = db.query(Document).filter(Document.id == document_id).first()
                failed_filename  = failed_doc.filename  if failed_doc else "Unknown"
                failed_author    = str(failed_doc.created_by) if failed_doc and failed_doc.created_by else "Unknown"

                asyncio.create_task(asyncio.to_thread(
                    webhook_service.trigger_webhook_background,
                    "document.failed",
                    {"document_id": document_id, "filename": failed_filename, "error": error_str},
                    failed_author,
                    SessionLocal
                ))
        finally:
            db.close()

    # ─────────────────────────────────────────────────────────────────────────
    # Read / list / stats — unchanged from original
    # ─────────────────────────────────────────────────────────────────────────

    @staticmethod
    def get_user_documents(
        db: Session,
        current_user,
        skip: int = 0,
        limit: int = 25,
        status_code=None,
        date_from=None,
        date_to=None,
        search_query=None,
        form_filters=None,
        document_type_id=None,
        client_id=None,
        uploaded_by=None,
        organisation_id=None,
        shared_only=False,
    ):
        from datetime import datetime, timedelta
        from sqlalchemy import and_, func

        base_query = DocumentService._document_access_query(db, current_user)

        if shared_only and isinstance(current_user, User):
            query = (
                db.query(Document)
                .join(DocumentShare, DocumentShare.document_id == Document.id)
                .filter(DocumentShare.user_id == current_user.id)
            )
        else:
            query = base_query

        query = query.options(joinedload(Document.status))
        query = query.options(joinedload(Document.form_data_relation))

        org_id = getattr(current_user, "context_organisation_id", None) or getattr(current_user, "organisation_id", None)
        if not org_id and hasattr(current_user, "id") and not hasattr(current_user, "organisation_id"):
            org_id = current_user.id
        if org_id:
            query = query.filter(Document.organisation_id == org_id)

        if uploaded_by:
            query = query.filter(Document.created_by == uploaded_by)

        if client_id:
            query = query.filter(Document.client_id == client_id)

        if status_code:
            code = str(status_code).upper()

            if code == "ARCHIVED":
                query = query.filter(Document.is_archived.is_(True))

            elif code == "PROCESSING":
                processing_codes = ["UPLOADING", "AI_QUEUED", "ANALYZING", "PROCESSING"]

                status_ids = (
                    db.query(Status.id)
                    .filter(Status.code.in_(processing_codes))
                    .all()
                )

                status_ids = [s.id for s in status_ids]

                query = query.filter(Document.status_id.in_(status_ids))
                query = query.filter(
                    or_(Document.is_archived == False, Document.is_archived.is_(None))
                )

            else:
                status = db.query(Status).filter(Status.code == code).first()
                if status:
                    query = query.filter(Document.status_id == status.id)

                query = query.filter(
                    or_(Document.is_archived == False, Document.is_archived.is_(None))
                )

        if document_type_id:
            query = query.filter(Document.document_type_id == document_type_id)

        if search_query:
            query = query.filter(Document.original_filename.ilike(f"%{search_query}%"))

        if date_from:
            try:
                query = query.filter(Document.created_at >= datetime.fromisoformat(date_from))
            except Exception:
                pass

        if date_to:
            try:
                query = query.filter(Document.created_at <= datetime.fromisoformat(date_to))
            except Exception:
                pass

        if form_filters:
            for field_id, value in form_filters.items():
                if not value:
                    continue
                query = query.join(DocumentFormData, DocumentFormData.document_id == Document.id)
                if "T" in str(value):
                    try:
                        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
                        start = dt.replace(hour=0, minute=0, second=0, microsecond=0)
                        end = start + timedelta(days=1)
                        query = query.filter(
                            and_(
                                func.nullif(DocumentFormData.data[field_id].astext, "") != None,
                                cast(func.nullif(DocumentFormData.data[field_id].astext, ""), DateTime) >= start,
                                cast(func.nullif(DocumentFormData.data[field_id].astext, ""), DateTime) < end,
                            )
                        )
                    except Exception:
                        pass
                else:
                    query = query.filter(DocumentFormData.data[field_id].astext == str(value))

        total = query.count()
        documents = query.order_by(Document.created_at.desc()).offset(skip).limit(limit).all()

        if not documents:
            return [], total

        doc_ids  = [d.id for d in documents]
        user_ids = {d.created_by for d in documents if d.created_by}
        org_ids  = {d.organisation_id for d in documents if d.organisation_id}

        users_map    = {u.id: u for u in db.query(User).filter(User.id.in_(user_ids)).all()}
        org_map      = {o.id: o for o in db.query(Organisation).filter(Organisation.id.in_(org_ids)).all()}
        form_rows    = db.query(DocumentFormData.document_id, DocumentFormData.data).filter(DocumentFormData.document_id.in_(doc_ids)).all()
        form_map     = {d_id: data for d_id, data in form_rows}
        fields       = db.query(FormField).all()
        field_map    = {str(f.id): f for f in fields}
        clients      = db.query(Client).all()
        client_map   = {str(c.id): c for c in clients}
        doc_types    = db.query(DocumentType).all()
        doc_type_map = {str(d.id): d for d in doc_types}

        result = []
        for doc in documents:
            raw = form_map.get(doc.id, {}) or {}
            client_name = None
            doc_type_name = None

            # ── Resolve client name ──────────────────────────────────────────
            # Priority 1: Document.client_id (set on the Document row directly)
            if doc.client_id:
                c = client_map.get(str(doc.client_id))
                if c:
                    client_name = c.business_name or f"{c.first_name} {c.last_name}"
            if not client_name and raw.get("client_id"):
                c = client_map.get(str(raw["client_id"]))
                if c:
                    client_name = c.business_name or f"{c.first_name} {c.last_name}"
            for field_id, value in raw.items():
                if field_id == "client_id":  
                    continue
                field = field_map.get(str(field_id))
                if not field or not value:
                    continue
                label = field.label.lower()
                if label == "client" and not client_name:
                    c = client_map.get(str(value))
                    if c:
                        client_name = c.business_name or f"{c.first_name} {c.last_name}"
                elif label == "document type" and not doc_type_name:
                    dt = doc_type_map.get(str(value))
                    if dt:
                        doc_type_name = dt.name

            uploaded_by_name = None
            if doc.created_by and doc.created_by in users_map:
                u = users_map[doc.created_by]
                uploaded_by_name = f"{u.first_name} {u.last_name}"
            elif doc.organisation_id and doc.organisation_id in org_map:
                uploaded_by_name = org_map[doc.organisation_id].name

            row = {
                "id": doc.id,
                "filename": doc.filename,
                "original_filename": doc.original_filename,
                "statusCode": doc.status.code if doc.status else None,
                "file_size": doc.file_size,
                "upload_progress": doc.upload_progress,
                "error_message": doc.error_message,
                "total_pages": doc.total_pages,
                "created_at": doc.created_at.isoformat() if doc.created_at else None,
                "updated_at": doc.updated_at.isoformat() if doc.updated_at else None,
                "is_archived": doc.is_archived,
                "uploaded_by": uploaded_by_name,
                "client": client_name,
                "document_type": doc_type_name,
                "form_data": copy.deepcopy(raw)
            }
            if doc.organisation_id in org_map:
                row["organisation_name"] = org_map[doc.organisation_id].name

            result.append(row)

        return result, total

    @staticmethod
    def _resolve_org_id(current_user):
        if not current_user:
            return None
        if getattr(current_user, "is_superuser", False):
            return None
        if isinstance(current_user, User):
            return str(current_user.organisation_id) if current_user.organisation_id else None
        return None

    @staticmethod
    def get_document_detail(db: Session, document_id: int, user: User):
        return (
            DocumentService._document_access_query(db, user)
            .options(joinedload(Document.extracted_documents))
            .options(joinedload(Document.unverified_documents))
            .filter(Document.id == document_id)
            .first()
        )

    @staticmethod
    def get_uploaded_by_filter(db: Session, current_user):
        query = db.query(
            User.id,
            func.concat(User.first_name, " ", User.last_name).label("name")
        )
        if getattr(current_user, "is_superuser", False):
            return query.all()
        org_id = DocumentService._resolve_org_id(current_user)
        if isinstance(current_user, User) and current_user.organisation_id:
            return query.filter(User.organisation_id == current_user.organisation_id).all()
        if isinstance(current_user, User) and current_user.client_id:
            return query.filter(User.client_id == current_user.client_id).all()
        return []

    @staticmethod
    async def archive_document(db: Session, document_id: int, user: User):
        document = DocumentService._get_accessible_document(db, document_id, user)
        if not document:
            return False
        document.is_archived = True
        db.commit()
        await websocket_manager.broadcast_document_status(
            document_id=document_id, status="ARCHIVED",
            user_id=str(document.created_by), progress=100
        )
        return True

    @staticmethod
    async def unarchive_document(db: Session, document_id: int, user: User):
        document = DocumentService._get_accessible_document(db, document_id, user)
        if not document:
            return False
        document.is_archived = False
        db.commit()
        await websocket_manager.broadcast_document_status(
            document_id=document_id, status="UNARCHIVED",
            user_id=str(document.created_by), progress=100
        )
        return True

    @staticmethod
    async def delete_document(db: Session, document_id: int, user: User):
        document = DocumentService._get_accessible_document(db, document_id, user)
        if not document:
            return None
        filename = document.original_filename or document.filename
        if document.s3_key:
            await s3_service.delete_file(document.s3_key)
        if document.analysis_report_s3_key:
            await s3_service.delete_file(document.analysis_report_s3_key)
        db.delete(document)
        db.commit()
        return filename

    @staticmethod
    async def cancel_document_analysis(db: Session, document_id: int, user: User):
        document = DocumentService._get_accessible_document(db, document_id, user)
        if not document:
            return False
        await DocumentService.update_document_status(
            db, document_id, "CANCELLED", error_message="Cancelled by user"
        )
        return True

    @staticmethod
    async def reanalyze_document(db: Session, document_id: int, user: User):
        document = DocumentService._get_accessible_document(db, document_id, user)
        if not document:
            raise Exception("Not allowed")
        if not document.s3_key:
            raise Exception("No file found")
        await DocumentService.update_document_status(
            db, document_id, "AI_QUEUED", progress=0, error_message="Reanalysis queued"
        )
        asyncio.create_task(
            DocumentService._perform_reanalysis_background(document_id, user.id)
        )
        return True

    @staticmethod
    async def _perform_reanalysis_background(document_id: int, created_by: str):
        import io as _io
        db = SessionLocal()
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document or not document.s3_key:
                return
            try:
                file_bytes = await s3_service.download_file(document.s3_key)
                file_content = _io.BytesIO(file_bytes)
            except Exception as e:
                await DocumentService.update_document_status(
                    db, document_id, "AI_FAILED",
                    error_message=f"Failed to retrieve file: {str(e)}"
                )
                return

            file_data = {
                'buffer': file_content,
                'filename': document.filename,
                'content_type': document.content_type
            }
            doc_type_id = str(document.document_type_id) if document.document_type_id else None
            template_id = str(document.template_id) if document.template_id else None

            await DocumentService._process_single_ai_analysis(
                document_id, file_data, doc_type_id, template_id
            )
        except Exception as e:
            print(f"Background re-analysis failed: {e}")
        finally:
            db.close()

    @staticmethod
    def get_document_stats(db: Session, user):
        base = DocumentService._document_access_query(db, user)
        total_all      = base.distinct(Document.id).count()
        total_archived = base.filter(Document.is_archived == True).count()

        status_counts = (
            base.join(Status, Status.id == Document.status_id)
            .filter(or_(Document.is_archived == False, Document.is_archived.is_(None)))
            .with_entities(Status.code, func.count(Document.id))
            .group_by(Status.code)
            .all()
        )
        counts = {code: count for code, count in status_counts}

        shared_with_me = 0
        if isinstance(user, User):
            shared_with_me = (
                db.query(Document)
                .join(DocumentShare, DocumentShare.document_id == Document.id)
                .filter(DocumentShare.user_id == user.id)
                .count()
            )

        return {
            "total": total_all,
            "processed": counts.get("COMPLETED", 0),
            "processing": (
                counts.get("PROCESSING", 0)
                + counts.get("ANALYZING", 0)
                + counts.get("AI_QUEUED", 0)
                + counts.get("UPLOADING", 0)
            ),
            "sharedWithMe": shared_with_me,
            "archived": total_archived,
        }


document_service = DocumentService()