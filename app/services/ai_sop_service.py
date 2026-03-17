# import asyncio
# import json
# import os
# import base64
# from io import BytesIO
# from zipfile import ZipFile
# from fastapi import HTTPException
# from pdfminer.high_level import extract_text
# from docx import Document
# from PIL import Image
# import pandas as pd
# from app.models.sop import SOP
# from app.models.status import Status
# from app.services.ai_client import openai_client
# from openpyxl import load_workbook
# from pdfminer.high_level import extract_text as pdf_extract

# # ── Shared JSON schema for table extraction ──────────────────────────────────
# PASS2_SCHEMA = """
# {
#   "coding_rules_cpt": [
#     {
#       "cptCode": null,
#       "description": null,
#       "ndcCode": null,
#       "units": null,
#       "chargePerUnit": null,
#       "modifier": null,
#       "replacementCPT": null
#     }
#   ],
#   "coding_rules_icd": [
#     {
#       "icdCode": null,
#       "description": null,
#       "notes": null
#     }
#   ]
# }
# """

# # ── Shared call helper ────────────────────────────────────────────────────────

# async def _call_ai(prompt: str, max_tokens: int = 8000) -> dict:
#     """
#     Single shared AI caller.
#     - Focused prompts need far fewer tokens than the old monolith
#     - 3-tier JSON repair fallback
#     """
#     response = await openai_client.chat.completions.create(
#         # model="gpt-4o-mini",
#         model="gpt-4o",  # gpt-4o-mini drops quality significantly on structured extraction tasks
#         messages=[
#             {
#                 "role": "system",
#                 "content": (
#                     "You are a medical SOP data extraction assistant. "
#                     "Always respond with valid JSON only. No markdown, no explanations."
#                 ),
#             },
#             {"role": "user", "content": prompt},
#         ],
#         temperature=0,
#         max_tokens=max_tokens,
#     )

#     raw = response.choices[0].message.content.strip()
#     if raw.startswith("```"):
#         raw = raw.replace("```json", "").replace("```", "").strip()

#     start = raw.find("{")
#     end = raw.rfind("}")

#     if start == -1:
#         raise HTTPException(422, f"AI returned non-JSON:\n{raw[:300]}")

#     raw = raw[start:] if (end == -1 or end < start) else raw[start : end + 1]

#     try:
#         return json.loads(raw)
#     except json.JSONDecodeError:
#         pass

#     try:
#         from json_repair import repair_json
#         return json.loads(repair_json(raw))
#     except Exception:
#         pass

#     last = raw.rfind("}")
#     if last > 0:
#         try:
#             return json.loads(raw[:last + 1])
#         except Exception:
#             pass

#     raise HTTPException(422, f"Cannot parse AI JSON:\n{raw[:400]}")


# # ── Vision helpers ─────────────────────────────────────────────────────────────

# def _is_garbled_text(text: str) -> bool:
#     """
#     Detect when pdfminer output is corrupted/unreadable.

#     Two signals:
#       1. >20% of non-whitespace chars are non-ASCII or bullet characters
#          (custom font encoding → each character extracted as its raw glyph ID)
#       2. Average word length < 2.5 (individual letters extracted as separate words)

#     Real SOP text scores: ~0.3% non-ASCII, avg word length ~4.7
#     Broken PDF scores:    ~27% non-ASCII, avg word length ~2.3
#     """
#     chars = [c for c in text if not c.isspace()]
#     if not chars:
#         return True  # empty → treat as garbled

#     non_ascii_ratio = sum(1 for c in chars if ord(c) > 127 or c == "•") / len(chars)
#     if non_ascii_ratio > 0.20:
#         return True

#     words = [w for w in text.split() if w]
#     if not words:
#         return True
#     avg_word_len = sum(len(w) for w in words[:1000]) / min(len(words), 1000)
#     if avg_word_len < 2.5:
#         return True

#     return False


# def _encode_pil(image: Image.Image, dpi_hint: int = 120, quality: int = 82) -> str:
#     """Convert a PIL image to a base64 JPEG string for the vision API."""
#     buf = BytesIO()
#     image.convert("RGB").save(buf, format="JPEG", quality=quality)
#     return base64.b64encode(buf.getvalue()).decode()


# async def _vision_text_from_images(
#     images_b64: list[str],
#     context: str = "medical SOP document",
# ) -> str:
#     """
#     Send a batch of base64-encoded page images to GPT-4o vision and return
#     extracted text.  Preserves table rows as pipe-delimited lines so that
#     _extract_table_sections() can still find them.
#     """
#     content: list[dict] = [
#         {
#             "type": "text",
#             "text": (
#                 f"You are an OCR engine for a {context}. "
#                 "Extract ALL visible text from the following page image(s) exactly as they appear. "
#                 "For tables: output each row as pipe-delimited text and wrap the table with "
#                 "--- TABLE N START --- / --- TABLE N END --- markers (increment N per table). "
#                 "Preserve section headings, bullet points, and all data values. "
#                 "Do NOT summarise or skip any content."
#             ),
#         }
#     ]
#     for i, b64 in enumerate(images_b64):
#         content.append({"type": "text", "text": f"--- Page {i + 1} ---"})
#         content.append({
#             "type": "image_url",
#             "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "high"},
#         })

#     response = await openai_client.chat.completions.create(
#         model="gpt-4o",          # must be gpt-4o for vision; gpt-4o-mini drops quality
#         messages=[{"role": "user", "content": content}],
#         temperature=0,
#         max_tokens=8000,
#     )
#     return response.choices[0].message.content.strip()


# async def _pdf_to_vision_text(path: str, dpi: int = 120, batch_size: int = 4) -> str:
#     """
#     Convert every page of a PDF to an image and extract text via GPT-4o vision.

#     Strategy:
#       - Convert pages at 120 dpi (≈164 KB/page as JPEG)
#       - Process in batches of `batch_size` pages per API call
#       - Run all batches concurrently (semaphore limits to 10 in-flight at once)
#       - Concatenate results in page order
#     """
#     from pdf2image import convert_from_path

#     loop = asyncio.get_event_loop()
#     pages: list[Image.Image] = await loop.run_in_executor(
#         None,
#         lambda: convert_from_path(path, dpi=dpi),
#     )

#     if not pages:
#         raise ValueError(f"pdf2image returned 0 pages for {path}")

#     print(f"[vision-pdf] {len(pages)} pages → batches of {batch_size}")

#     # encode all pages (CPU-bound, run in executor)
#     encoded: list[str] = await loop.run_in_executor(
#         None,
#         lambda: [_encode_pil(p, dpi) for p in pages],
#     )

#     # build batches
#     batches: list[list[str]] = [
#         encoded[i : i + batch_size] for i in range(0, len(encoded), batch_size)
#     ]

#     sem = asyncio.Semaphore(10)

#     async def _process_batch(batch_imgs: list[str], batch_no: int) -> tuple[int, str]:
#         async with sem:
#             text = await _vision_text_from_images(batch_imgs)
#             return batch_no, text

#     tasks = [_process_batch(batch, i) for i, batch in enumerate(batches)]
#     results: list[tuple[int, str]] = await asyncio.gather(*tasks)

#     # sort by batch index and join
#     results.sort(key=lambda x: x[0])
#     return "\n\n".join(r[1] for r in results)


# async def _docx_embedded_image_text(path: str) -> str:
#     """
#     Extract text from images embedded inside a DOCX file (word/media/*).
#     Returns concatenated vision-extracted text, or "" if none found.
#     """
#     try:
#         with ZipFile(path) as zf:
#             image_names = [
#                 n for n in zf.namelist()
#                 if n.startswith("word/media/")
#                 and n.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"))
#             ]

#         if not image_names:
#             return ""

#         print(f"[vision-docx] {len(image_names)} embedded image(s) found")

#         async def _extract_one(img_name: str) -> str:
#             with ZipFile(path) as zf:
#                 img_bytes = zf.read(img_name)

#             img = Image.open(BytesIO(img_bytes))
#             b64 = _encode_pil(img)
#             return await _vision_text_from_images([b64], context="medical SOP embedded image")

#         parts = await asyncio.gather(*[_extract_one(n) for n in image_names])
#         return "\n\n".join(p for p in parts if p.strip())

#     except Exception as e:
#         print(f"[vision-docx] embedded image extraction failed: {e}")
#         return ""


# # ── Focused parallel extractors ───────────────────────────────────────────────

# async def _extract_meta(text: str) -> dict:
#     """Extracts: basic_information + provider_information + workflow_process"""
#     prompt = f"""Extract ONLY these sections from the medical SOP below.

# Return ONLY this JSON (no extra keys):
# {{
#   "basic_information": {{"sop_title": "", "category": ""}},
#   "provider_information": {{
#     "billing_provider_name": "", "billing_provider_npi": "",
#     "provider_tax_id": "", "billing_address": "",
#     "software": "", "clearinghouse": ""
#   }},
#   "workflow_process": {{
#     "workflow_description": "",
#     "eligibility_verification_portals": [],
#     "posting_charges_rules": []
#   }}
# }}

# Rules:
# - Provider info may appear in headers, footers, letterhead, or signature blocks — search ENTIRE document.
# - eligibility_verification_portals: list portal names/URLs.
# - posting_charges_rules: list rules about how charges are posted.
# - Use "" for missing fields, NOT null.

# DOCUMENT:
# {text[:12000]}"""
#     return await _call_ai(prompt, max_tokens=2000)


# async def _extract_billing(text: str) -> dict:
#     """Extracts: billing_guidelines only"""
#     prompt = f"""Extract ONLY billing guidelines from the medical SOP below.

# Billing guidelines = operational/documentation rules ONLY.
# Examples: authorization requirements, claim submission rules, timely filing, documentation requirements.

# STRICT EXCLUSIONS — do NOT include rules that contain:
# - CPT codes (numeric like 99213, J0129)
# - ICD-10 codes (letter+number like M17.0)
# - Payer-specific rules (those belong in payer_guidelines)

# Group rules by their heading/category. Infer category name from surrounding text.

# Return ONLY this JSON:
# {{
#   "billing_guidelines": [
#     {{
#       "category": "<heading from document>",
#       "rules": [{{"description": "<exact original text>"}}]
#     }}
#   ]
# }}

# DOCUMENT:
# {text[:16000]}"""
#     return await _call_ai(prompt, max_tokens=4000)


# async def _extract_payers(text: str) -> dict:
#     """Extracts: payer_guidelines only"""
#     prompt = f"""Extract ONLY payer-specific guidelines from the medical SOP below.

# MUST extract:
# - ERA setup status (e.g. "Completed", "Form Submitted")
# - EDI setup info
# - Claim mailing addresses
# - Network/credentialing status (INN/OON/NA)
# - Timely Filing Limits (TFL)
# - Payer ID numbers
# - Any rule mentioning a specific payer (Medicare, Medicaid, Aetna, BCBS, UHC, Cigna, etc.)

# CRITICAL MERGING RULE:
# If the same payer appears in multiple tables/sections, merge ALL their data into ONE object.

# STRICT ASSOCIATION:
# Only link data (payerId, tfl, etc.) to a payer if EXPLICITLY stated for that payer. No proximity guessing.

# Return ONLY this JSON:
# {{
#   "payer_guidelines": [
#     {{
#       "payerName": "", "description": "", "payerId": "",
#       "eraStatus": "", "ediStatus": "", "tfl": "",
#       "networkStatus": "", "mailingAddress": ""
#     }}
#   ]
# }}

# DOCUMENT:
# {text[:16000]}"""
#     return await _call_ai(prompt, max_tokens=6000)


# async def _extract_coding(text: str) -> dict:
#     """
#     Extracts CPT rules from narrative text (inline modifier rules, replacement
#     rules, unit caps, admin code rules) AND all ICD rules (replace-with,
#     do-not-bill-together, use-only-when).
#     Table rows are handled separately by _extract_coding_from_tables.
#     """
#     prompt = f"""You are extracting CPT and ICD coding rules from a medical SOP.

# CPT RULES  ->  coding_rules_cpt
# Extract EVERY sentence or rule that mentions a CPT or HCPCS code.

# CPT/HCPCS code formats include:
#   Numeric: 99213, 73502, 73521, 77080, 96372, 96365, 96366, 96413, 81002
#   J-codes: J0129, J0897, J1010, J1200, J2507, J2919, J3489, J7321, j7050, etc.

# For EACH rule create ONE object:
#   cptCode       -> main CPT/HCPCS code mentioned (e.g. "J0129", "73502")
#   description   -> full rule text exactly as written in the document
#   modifier      -> any modifier mentioned (e.g. "JZ", "JA", "LT", "RT", "50", "59", "95", "EJ")
#   replacementCPT-> if rule says replace X with Y (e.g. "73521")
#   units         -> if a specific unit count is mentioned (e.g. "120 max")
#   ndcCode       -> leave blank (NDC codes are in tables, not narrative)
#   chargePerUnit -> leave blank (charges are in tables, not narrative)

# Examples:
#   "Please use JZ modifier in all J code CPTs"  -> cptCode:"J-codes", modifier:"JZ"
#   "JA modifier should be only use with Medicare for CPT J0129" -> cptCode:"J0129", modifier:"JA"
#   "CPT 73502 marked on super bill with LR & RT then use 73521" -> cptCode:"73502", replacementCPT:"73521"
#   "CPT J1010 maximum can be billed with 120 units only" -> cptCode:"J1010", units:"120 max"
#   "If J7321 is billed more than once, use modifier EJ" -> cptCode:"J7321", modifier:"EJ"
#   "For CPT J0897 use admin code 96372 with Aetna" -> TWO objects: J0897 and 96372

# ICD RULES  ->  coding_rules_icd
# Extract EVERY sentence or rule mentioning an ICD-10 diagnosis code.

# ICD-10 format: starts with a LETTER followed by numbers (M17.0, Z00.00, L93.0, M54.50, M1A9XX1)

# For EACH ICD code create ONE object:
#   icdCode     -> the ICD code exactly as written (e.g. "M17.0", "M54.50")
#   description -> full rule text exactly as written
#   notes       -> key instruction (e.g. "Replace with M54.59", "Do not bill with M05.89")

# CRITICAL: If one sentence mentions multiple ICD codes, create a SEPARATE object for each:
#   "M54.50 replace with M54.59"   -> icdCode:"M54.50", notes:"Replace with M54.59"
#   "M25.50 replace with M25.59"   -> icdCode:"M25.50", notes:"Replace with M25.59"
#   "Do not bill both M45.9 and M05.89 together" -> TWO objects: M45.9 AND M05.89
#   "Use M17.0 only when M17.12, M17.9 & M17.11 is given" -> icdCode:"M17.0"

# NEVER put CPT codes in coding_rules_icd.
# NEVER put ICD codes in coding_rules_cpt.

# Return ONLY valid JSON (no markdown, no extra text):
# {{
#   "coding_rules_cpt": [
#     {{"cptCode": "", "description": "", "ndcCode": "", "units": "", "chargePerUnit": "", "modifier": "", "replacementCPT": ""}}
#   ],
#   "coding_rules_icd": [
#     {{"icdCode": "", "description": "", "notes": ""}}
#   ]
# }}

# DOCUMENT:
# {text[:20000]}"""
#     return await _call_ai(prompt, max_tokens=12000)


# async def _extract_coding_from_tables(text: str) -> dict:
#     """
#     Extracts ALL CPT/HCPCS rows from structured tables.
#     Handles BOTH infusion/drug NDC tables AND x-ray modifier/replacement tables.
#     """
#     table_text = _extract_table_sections(text)
#     if not table_text.strip():
#         return {"coding_rules_cpt": [], "coding_rules_icd": []}

#     prompt = f"""Extract ALL CPT and HCPCS code rows from these medical document tables.

# There are TWO types of tables — handle both:

# TYPE 1 — Infusion/Drug table (columns: CPT Code | Description | NDC Code | Units | Charge per Unit)
#   For each row:
#   - cptCode      -> J-code or CPT (e.g. "J0129", "J1010")
#   - description  -> drug/procedure name (e.g. "ORENCIA (Injection, abatacept, 10Mg)")
#   - ndcCode      -> NDC code exactly as shown (e.g. "00003-2187-13")
#   - units        -> unit count or "As per received superbill"
#   - chargePerUnit-> charge amount (e.g. "$100.00", "$3750.00")
#   - modifier     -> leave blank
#   - replacementCPT -> leave blank

# TYPE 2 — X-ray/Modifier table (columns: CPT Code | Modifier | Replacement CPT Code)
#   For each row:
#   - cptCode      -> the CPT code (e.g. "73600", "73510", "72040")
#   - description  -> "X-ray code"
#   - modifier     -> modifier(s) listed (e.g. "50/LT/RT", "NO Modifier", "Deleted code")
#   - replacementCPT -> replacement code if listed (e.g. "73521"), else blank
#   - ndcCode, units, chargePerUnit -> leave blank

# Rules:
# - Extract EVERY row — do NOT skip any
# - Preserve exact values (NDC codes, dollar amounts, unit counts)
# - Do NOT hallucinate rows not present in the table

# Return ONLY valid JSON (no markdown):
# {{
#   "coding_rules_cpt": [
#     {{"cptCode": "", "description": "", "ndcCode": "", "units": "", "chargePerUnit": "", "modifier": "", "replacementCPT": ""}}
#   ],
#   "coding_rules_icd": []
# }}

# TABLES:
# {table_text}"""
#     return await _call_ai(prompt, max_tokens=8000)


# def _extract_table_sections(text: str) -> str:
#     tables = []
#     lines = text.splitlines()
#     capture = False
#     current: list = []
#     for line in lines:
#         if "TABLE" in line and "START" in line:
#             capture = True
#             current = []
#         elif "TABLE" in line and "END" in line:
#             capture = False
#             tables.append("\n".join(current))
#         elif capture:
#             current.append(line)
#     return "\n\n".join(tables)


# class AISOPService:

#     # ── Text extraction ──────────────────────────────────────────────────────

#     @staticmethod
#     def extract_pdf_text(path: str) -> str:
#         return extract_text(path)

#     @staticmethod
#     def process_sop_extraction(
#         sop_id: str,
#         file_path: str = None,
#         content_type: str = None,
#         s3_key: str = None,
#     ):
#         from app.core.database import SessionLocal
#         from app.services.s3_service import s3_service
#         import tempfile

#         db = SessionLocal()
#         temp_file_to_remove = None

#         try:
#             if s3_key and not file_path:
#                 file_extension = s3_key.split(".")[-1] if "." in s3_key else ""
#                 with tempfile.NamedTemporaryFile(
#                     delete=False,
#                     suffix=f".{file_extension}" if file_extension else "",
#                 ) as tmp:
#                     file_data = asyncio.run(s3_service.download_file(s3_key))
#                     tmp.write(file_data)
#                     file_path = tmp.name
#                     temp_file_to_remove = file_path

#             if not file_path:
#                 raise Exception("No file source provided for extraction")

#             text = asyncio.run(AISOPService.extract_text(file_path))
#             structured = asyncio.run(AISOPService.ai_extract_sop_structured(text))

#             active_status = db.query(Status).filter(
#                 Status.code == "ACTIVE", Status.type == "GENERAL"
#             ).first()

#             sop = db.query(SOP).filter(SOP.id == sop_id).first()
#             if not sop:
#                 return

#             structured = AISOPService.normalize_ai_sop(structured)

#             sop.title = (
#                 structured.get("basic_information", {}).get("sop_title") or sop.title
#             )
#             sop.category = (
#                 structured.get("basic_information", {}).get("category") or sop.category
#             )
#             provider_info = structured.get("provider_information") or {}

#             if not provider_info.get("billingProviderNPI") and sop.client:
#                 provider_info["billingProviderNPI"] = sop.client.npi

#             sop.provider_info = provider_info
#             sop.workflow_process = structured.get("workflow_process")

#             from app.models.sop import SOPDocument

#             doc = db.query(SOPDocument).filter(
#                 SOPDocument.sop_id == sop_id,
#                 SOPDocument.s3_key == s3_key if s3_key else None,
#             ).first()

#             if not doc:
#                 doc = db.query(SOPDocument).filter(
#                     SOPDocument.sop_id == sop_id,
#                     SOPDocument.category == "Source file",
#                 ).first()

#             if doc:
#                 source_name = "source_file" if doc.category == "Source file" else (doc.name or "Document")

#                 def inject_source(items):
#                     if not items:
#                         return items
#                     for item in items:
#                         if isinstance(item, dict):
#                             if "rules" in item:
#                                 for r in item["rules"]:
#                                     r["source"] = source_name
#                             else:
#                                 item["source"] = source_name
#                     return items

#                 if doc.category == "Source file":
#                     doc.billing_guidelines = structured.get("billing_guidelines")
#                     doc.payer_guidelines = structured.get("payer_guidelines")
#                     doc.coding_rules_cpt = structured.get("coding_rules_cpt")
#                     doc.coding_rules_icd = structured.get("coding_rules_icd")
#                 else:
#                     doc.billing_guidelines = inject_source(structured.get("billing_guidelines"))
#                     doc.payer_guidelines = inject_source(structured.get("payer_guidelines"))
#                     doc.coding_rules_cpt = inject_source(structured.get("coding_rules_cpt"))
#                     doc.coding_rules_icd = inject_source(structured.get("coding_rules_icd"))

#                 doc.processed = True

#             sop.status_id = active_status.id if active_status else sop.status_id
#             db.commit()

#         except Exception as e:
#             print(f"Error in background extraction: {e}")
#             failed_status = db.query(Status).filter(
#                 Status.code == "FAILED", Status.type == "DOCUMENT"
#             ).first()
#             sop = db.query(SOP).filter(SOP.id == sop_id).first()
#             if sop and failed_status:
#                 sop.status_id = failed_status.id
#                 db.commit()

#         finally:
#             db.close()
#             target_path = temp_file_to_remove or file_path
#             if target_path and os.path.exists(target_path):
#                 try:
#                     os.remove(target_path)
#                 except Exception:
#                     pass

#     @staticmethod
#     async def extract_image_text(path: str) -> str:
#         with Image.open(path) as img:
#             buffered = BytesIO()
#             img.convert("RGB").save(buffered, format="JPEG", quality=92)
#             image_bytes = buffered.getvalue()

#         image_base64 = base64.b64encode(image_bytes).decode("utf-8")

#         response = await openai_client.chat.completions.create(
#             # model="gpt-4o-mini",
#             model="gpt-4o",  # gpt-4o-mini drops quality significantly on OCR tasks
#             messages=[
#                 {
#                     "role": "system",
#                     "content": "You are an OCR engine. Extract ALL visible text exactly as-is. No summaries.",
#                 },
#                 {
#                     "role": "user",
#                     "content": [
#                         {
#                             "type": "text",
#                             "text": "Extract all text. Preserve line breaks.",
#                         },
#                         {
#                             "type": "image_url",
#                             "image_url": {
#                                 "url": f"data:image/jpeg;base64,{image_base64}",
#                                 "detail": "high",
#                             },
#                         },
#                     ],
#                 },
#             ],
#             temperature=0,
#             max_tokens=4000,
#         )

#         text = response.choices[0].message.content.strip()
#         if not text:
#             raise HTTPException(422, "No text extracted from image")
#         return text

#     @staticmethod
#     def extract_docx_text(path: str) -> str:
#         doc = Document(path)
#         parts = []
#         for p in doc.paragraphs:
#             if p.text.strip():
#                 parts.append(p.text.strip())
#         for table_index, table in enumerate(doc.tables, start=1):
#             parts.append(f"\n--- TABLE {table_index} START ---")
#             for row in table.rows:
#                 cells = [cell.text.strip() for cell in row.cells]
#                 if any(cells):
#                     parts.append(" | ".join(cells))
#             parts.append(f"--- TABLE {table_index} END ---\n")
#         return "\n".join(parts)

#     @staticmethod
#     def extract_excel_text(path: str) -> str:
#         ext = os.path.splitext(path)[1].lower()
#         try:
#             if ext == ".xls":
#                 df_dict = pd.read_excel(path, sheet_name=None, engine="xlrd")
#             elif ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
#                 df_dict = pd.read_excel(path, sheet_name=None, engine="openpyxl")
#             else:
#                 raise ValueError(f"Unsupported Excel extension: {ext}")
#             full_text = ""
#             for sheet_name, df in df_dict.items():
#                 full_text += f"\n--- Sheet: {sheet_name} ---\n"
#                 full_text += df.fillna("").astype(str).to_string(index=False)
#             return full_text
#         except Exception as e:
#             raise Exception(f"Excel extraction failed: {str(e)}")

#     @staticmethod
#     async def extract_text(path: str) -> str:
#         ext = os.path.splitext(path)[1].lower()

#         # ── PDF ──────────────────────────────────────────────────────────────
#         if ext == ".pdf":
#             text = pdf_extract(path)

#             if _is_garbled_text(text):
#                 # Font-encoded / scanned PDF — pdfminer output is unusable.
#                 # Fall back to page-image vision extraction via GPT-4o.
#                 print(
#                     f"[extract_text] PDF text garbled (font encoding issue) — "
#                     f"switching to vision extraction for {os.path.basename(path)}"
#                 )
#                 text = await _pdf_to_vision_text(path)

#             return text

#         # ── DOCX ─────────────────────────────────────────────────────────────
#         if ext == ".docx":
#             doc = Document(path)
#             parts = []

#             for p in doc.paragraphs:
#                 if p.text.strip():
#                     parts.append(p.text.strip())

#             # Tables with TABLE markers so _extract_table_sections() can find them
#             for table_index, table in enumerate(doc.tables, start=1):
#                 parts.append(f"\n--- TABLE {table_index} START ---")
#                 for row in table.rows:
#                     row_text = " | ".join(
#                         cell.text.strip() for cell in row.cells if cell.text.strip()
#                     )
#                     if row_text:
#                         parts.append(row_text)
#                 parts.append(f"--- TABLE {table_index} END ---\n")

#             text = "\n".join(parts)

#             # Embedded images (superbill screenshots, scanned inserts, etc.)
#             embedded_text = await _docx_embedded_image_text(path)
#             if embedded_text:
#                 print(
#                     f"[extract_text] Appending vision text from "
#                     f"{len(embedded_text.split(chr(10)))} lines of embedded images"
#                 )
#                 text = text + "\n\n--- EMBEDDED IMAGE CONTENT ---\n" + embedded_text

#             return text

#         # ── Excel ─────────────────────────────────────────────────────────────
#         if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
#             wb = load_workbook(path, data_only=True)
#             parts = []
#             for sheet in wb.worksheets:
#                 parts.append(f"\n--- SHEET: {sheet.title} ---")
#                 for row in sheet.iter_rows(values_only=True):
#                     row_text = " | ".join(
#                         str(cell) for cell in row if cell is not None
#                     )
#                     if row_text.strip():
#                         parts.append(row_text)
#             return "\n".join(parts)

#         if ext == ".xls":
#             df_dict = pd.read_excel(path, sheet_name=None, engine="xlrd")
#             parts = []
#             for sheet_name, df in df_dict.items():
#                 parts.append(f"\n--- SHEET: {sheet_name} ---")
#                 parts.append(df.fillna("").astype(str).to_string(index=False))
#             return "\n".join(parts)

#         # ── Images ────────────────────────────────────────────────────────────
#         if ext in [".png", ".jpg", ".jpeg"]:
#             return await AISOPService.extract_image_text(path)

#         raise ValueError(f"Unsupported file type: {ext}")

#     @staticmethod
#     def normalize_ai_sop(data: dict) -> dict:
#         category = (
#             data.get("basic_information", {}).get("category")
#             if isinstance(data.get("basic_information"), dict)
#             else data.get("category")
#         )
#         if isinstance(category, dict):
#             category = category.get("title", "")

#         if "basic_information" not in data:
#             data["basic_information"] = {}

#         data["basic_information"]["category"] = category or data.get("category", "")
#         data["category"] = data["basic_information"]["category"]

#         source = "AI Extracted"

#         provider_info_raw = data.get("provider_information") or {}
#         data["provider_information"] = {
#             "providerName": provider_info_raw.get("billing_provider_name")
#             or provider_info_raw.get("providerName"),
#             "billingProviderName": provider_info_raw.get("billing_provider_name")
#             or provider_info_raw.get("billingProviderName"),
#             "billingProviderNPI": provider_info_raw.get("billing_provider_npi")
#             or provider_info_raw.get("billingProviderNPI"),
#             "providerTaxID": provider_info_raw.get("provider_tax_id")
#             or provider_info_raw.get("providerTaxID"),
#             "billingAddress": provider_info_raw.get("billing_address")
#             or provider_info_raw.get("billingAddress"),
#             "software": provider_info_raw.get("software"),
#             "clearinghouse": provider_info_raw.get("clearinghouse"),
#         }

#         workflow_raw = data.get("workflow_process") or {}
#         data["workflow_process"] = {
#             "description": workflow_raw.get("workflow_description")
#             or workflow_raw.get("description"),
#             "eligibility_verification_portals": workflow_raw.get(
#                 "eligibility_verification_portals", []
#             ),
#             "posting_charges_rules": workflow_raw.get("posting_charges_rules", []),
#         }

#         guidelines = data.get("billing_guidelines", [])
#         normalized_guidelines = []
#         for group in guidelines:
#             if not isinstance(group, dict):
#                 continue
#             normalized_guidelines.append(
#                 {
#                     "category": group.get("category", "Guidelines"),
#                     "rules": [
#                         {"description": r.get("description", ""), "source": source}
#                         for r in group.get("rules", [])
#                         if isinstance(r, dict)
#                     ],
#                 }
#             )
#         data["billing_guidelines"] = normalized_guidelines

#         payer_guidelines = data.get("payer_guidelines", [])
#         normalized_payers = []
#         for pg in payer_guidelines:
#             if isinstance(pg, str):
#                 normalized_payers.append(
#                     {
#                         "payerName": "Unknown",
#                         "description": pg,
#                         "payerId": "",
#                         "eraStatus": "",
#                         "tfl": "",
#                         "networkStatus": "",
#                         "mailingAddress": "",
#                         "source": source,
#                     }
#                 )
#             elif isinstance(pg, dict):
#                 normalized_payers.append(
#                     {
#                         "payerName": pg.get("payerName")
#                         or pg.get("payer_name")
#                         or pg.get("title")
#                         or pg.get("payer")
#                         or "Unknown",
#                         "description": pg.get("description") or "",
#                         "payerId": pg.get("payerId") or pg.get("payer_id") or "",
#                         "eraStatus": pg.get("eraStatus") or pg.get("era_status") or "",
#                         "ediStatus": pg.get("ediStatus") or pg.get("edi_status") or "",
#                         "tfl": pg.get("tfl") or "",
#                         "networkStatus": pg.get("networkStatus")
#                         or pg.get("network_status")
#                         or "",
#                         "mailingAddress": pg.get("mailingAddress")
#                         or pg.get("mailing_address")
#                         or pg.get("address")
#                         or "",
#                         "source": pg.get("source") or source,
#                     }
#                 )
#         data["payer_guidelines"] = normalized_payers

#         for key in ("coding_rules_cpt", "coding_rules_icd"):
#             data[key] = [
#                 {**r, "source": source}
#                 for r in data.get(key, [])
#                 if isinstance(r, dict)
#             ]

#         return data

#     @staticmethod
#     def safe_parse_json(raw: str) -> dict:
#         raw = raw.strip()
#         if raw.startswith("```"):
#             raw = raw.replace("```json", "").replace("```", "").strip()
#         if not raw.startswith("{"):
#             raise ValueError("AI did not return JSON")
#         try:
#             return json.loads(raw)
#         except json.JSONDecodeError:
#             last_brace = raw.rfind("}")
#             if last_brace != -1:
#                 try:
#                     return json.loads(raw[: last_brace + 1])
#                 except ValueError as e:
#                     raise HTTPException(422, str(e))
#                 except Exception:
#                     raise HTTPException(500, "Internal server error")
#             raise ValueError(f"Invalid JSON from AI:\n{raw[:1000]}")

#     @staticmethod
#     def extract_table_sections(text: str) -> str:
#         return _extract_table_sections(text)

#     # ── Legacy compat ────────────────────────────────────────────────────────

#     @staticmethod
#     async def _call_ai(prompt: str) -> dict:
#         return await _call_ai(prompt)

#     @staticmethod
#     async def extract_narrative_and_rules(text: str) -> dict:
#         meta, billing, payers, coding = await asyncio.gather(
#             _extract_meta(text),
#             _extract_billing(text),
#             _extract_payers(text),
#             _extract_coding(text),
#         )
#         return {**meta, **billing, **payers, **coding}

#     @staticmethod
#     async def extract_tables(text: str) -> dict:
#         return await _extract_coding_from_tables(text)

#     # ── Main entry point ─────────────────────────────────────────────────────

#     @staticmethod
#     async def ai_extract_sop_structured(text: str) -> dict:
#         """
#         Fires all 5 extraction calls simultaneously via asyncio.gather.
#         Total time ≈ slowest single call (~5-10s) instead of sum of all calls.

#         Calls:
#           1. _extract_meta        → title, provider, workflow
#           2. _extract_billing     → billing guidelines
#           3. _extract_payers      → payer guidelines
#           4. _extract_coding      → CPT + ICD from narrative text
#           5. _extract_coding_from_tables → CPT from table sections
#         """
#         (
#             meta,
#             billing,
#             payers,
#             coding_narrative,
#             coding_tables,
#         ) = await asyncio.gather(
#             _extract_meta(text),
#             _extract_billing(text),
#             _extract_payers(text),
#             _extract_coding(text),
#             _extract_coding_from_tables(text),
#         )

#         # Merge + deduplicate CPT from narrative and table passes
#         cpt_combined: list = (
#             coding_narrative.get("coding_rules_cpt", [])
#             + coding_tables.get("coding_rules_cpt", [])
#         )
#         seen_cpt: set = set()
#         deduped_cpt = []
#         for r in cpt_combined:
#             key = r.get("cptCode") or r.get("description", "")
#             if key and key not in seen_cpt:
#                 seen_cpt.add(key)
#                 deduped_cpt.append(r)
#             elif not key:
#                 deduped_cpt.append(r)

#         final = {
#             **meta,
#             **billing,
#             **payers,
#             "coding_rules_cpt": deduped_cpt,
#             "coding_rules_icd": coding_narrative.get("coding_rules_icd", []),
#         }

#         return AISOPService.normalize_ai_sop(final)

import asyncio
import json
import os
import base64
from io import BytesIO
from zipfile import ZipFile
from fastapi import HTTPException
from pdfminer.high_level import extract_text
from docx import Document
from PIL import Image
import pandas as pd
from phonenumbers import data
from app.models.sop import SOP
from app.models.status import Status
from app.services.ai_client import openai_client
from openpyxl import load_workbook
from pdfminer.high_level import extract_text as pdf_extract

import re as _re_module

def _repair_j_code(code: str) -> str:
    """
    Repair J-codes where pdfminer dropped the leading 'J' character due to
    PDF font encoding issues. This affects pages where J0XXX/JXXXX codes
    appear as just their numeric part (e.g. "702" instead of "J0702").
    
    Rules:
      3 digits  → "J0" + code  (702 → J0702, 897 → J0897)
      4 digits  → "J"  + code  (1010 → J1010, 7321 → J7321)
      Anything else → unchanged
    """
    if not code:
        return code
    c = code.strip()
    # Already a valid CPT/HCPCS — don't touch
    if _re_module.match(r'^[A-Z][0-9]{4}[A-Z0-9]?$', c):
        return c
    if _re_module.match(r'^[0-9]{5}[A-Z]?$', c):
        return c  # 5-digit CPT
    # 3 bare digits → J0XXX
    if _re_module.match(r'^[0-9]{3}$', c):
        return f"J0{c}"
    # 4 bare digits → JXXXX
    if _re_module.match(r'^[0-9]{4}$', c):
        return f"J{c}"
    return c


# ── Shared JSON schema for table extraction ──────────────────────────────────
PASS2_SCHEMA = """
{
  "coding_rules_cpt": [
    {
      "cptCode": null,
      "description": null,
      "ndcCode": null,
      "units": null,
      "chargePerUnit": null,
      "modifier": null,
      "replacementCPT": null
    }
  ],
  "coding_rules_icd": [
    {
      "icdCode": null,
      "description": null,
      "notes": null
    }
  ]
}
"""

# ── Shared call helper ────────────────────────────────────────────────────────

async def _call_ai(prompt: str, max_tokens: int = 8000) -> dict:
    """
    Single shared AI caller.
    - Focused prompts need far fewer tokens than the old monolith
    - 3-tier JSON repair fallback
    """
    response = await openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a medical SOP data extraction assistant. "
                    "Always respond with valid JSON only. No markdown, no explanations."
                ),
            },
            {"role": "user", "content": prompt},
        ],
        temperature=0,
        max_tokens=max_tokens,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()

    start = raw.find("{")
    end = raw.rfind("}")

    if start == -1:
        raise HTTPException(422, f"AI returned non-JSON:\n{raw[:300]}")

    raw = raw[start:] if (end == -1 or end < start) else raw[start : end + 1]

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    try:
        from json_repair import repair_json
        return json.loads(repair_json(raw))
    except Exception:
        pass

    last = raw.rfind("}")
    if last > 0:
        try:
            return json.loads(raw[:last + 1])
        except Exception:
            pass

    raise HTTPException(422, f"Cannot parse AI JSON:\n{raw[:400]}")


# ── Vision helpers ─────────────────────────────────────────────────────────────

def _is_garbled_text(text: str) -> bool:
    """
    Detect when pdfminer output is corrupted/unreadable.

    Two signals:
      1. >20% of non-whitespace chars are non-ASCII or bullet characters
         (custom font encoding → each character extracted as its raw glyph ID)
      2. Average word length < 2.5 (individual letters extracted as separate words)

    Real SOP text scores: ~0.3% non-ASCII, avg word length ~4.7
    Broken PDF scores:    ~27% non-ASCII, avg word length ~2.3
    """
    chars = [c for c in text if not c.isspace()]
    if not chars:
        return True  # empty → treat as garbled

    non_ascii_ratio = sum(1 for c in chars if ord(c) > 127 or c == "•") / len(chars)
    if non_ascii_ratio > 0.20:
        return True

    words = [w for w in text.split() if w]
    if not words:
        return True
    avg_word_len = sum(len(w) for w in words[:1000]) / min(len(words), 1000)
    if avg_word_len < 2.5:
        return True

    return False


def _encode_pil(image: Image.Image, dpi_hint: int = 120, quality: int = 82) -> str:
    """Convert a PIL image to a base64 JPEG string for the vision API."""
    buf = BytesIO()
    image.convert("RGB").save(buf, format="JPEG", quality=quality)
    return base64.b64encode(buf.getvalue()).decode()


async def _vision_text_from_images(
    images_b64: list[str],
    context: str = "medical SOP document",
) -> str:
    """
    Send a batch of base64-encoded page images to GPT-4o vision and return
    extracted text.  Preserves table rows as pipe-delimited lines so that
    _extract_table_sections() can still find them.
    """
    content: list[dict] = [
        {
            "type": "text",
            "text": (
                f"You are an OCR engine for a {context}. "
                "Extract ALL visible text from the following page image(s) exactly as they appear. "
                "For tables: output each row as pipe-delimited text and wrap the table with "
                "--- TABLE N START --- / --- TABLE N END --- markers (increment N per table). "
                "Preserve section headings, bullet points, and all data values. "
                "Do NOT summarise or skip any content."
            ),
        }
    ]
    for i, b64 in enumerate(images_b64):
        content.append({"type": "text", "text": f"--- Page {i + 1} ---"})
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "high"},
        })

    response = await openai_client.chat.completions.create(
        model="gpt-4o",          # must be gpt-4o for vision; gpt-4o-mini drops quality
        messages=[{"role": "user", "content": content}],
        temperature=0,
        max_tokens=8000,
    )
    return response.choices[0].message.content.strip()


async def _pdf_to_vision_text(path: str, dpi: int = 120, batch_size: int = 4) -> str:
    """
    Convert every page of a PDF to an image and extract text via GPT-4o vision.

    Strategy:
      - Convert pages at 120 dpi (≈164 KB/page as JPEG)
      - Process in batches of `batch_size` pages per API call
      - Run all batches concurrently (semaphore limits to 10 in-flight at once)
      - Concatenate results in page order
    """
    from pdf2image import convert_from_path

    loop = asyncio.get_event_loop()
    pages: list[Image.Image] = await loop.run_in_executor(
        None,
        lambda: convert_from_path(path, dpi=dpi),
    )

    if not pages:
        raise ValueError(f"pdf2image returned 0 pages for {path}")

    print(f"[vision-pdf] {len(pages)} pages → batches of {batch_size}")

    # encode all pages (CPU-bound, run in executor)
    encoded: list[str] = await loop.run_in_executor(
        None,
        lambda: [_encode_pil(p, dpi) for p in pages],
    )

    # build batches
    batches: list[list[str]] = [
        encoded[i : i + batch_size] for i in range(0, len(encoded), batch_size)
    ]

    sem = asyncio.Semaphore(10)

    async def _process_batch(batch_imgs: list[str], batch_no: int) -> tuple[int, str]:
        async with sem:
            text = await _vision_text_from_images(batch_imgs)
            return batch_no, text

    tasks = [_process_batch(batch, i) for i, batch in enumerate(batches)]
    results: list[tuple[int, str]] = await asyncio.gather(*tasks)

    # sort by batch index and join
    results.sort(key=lambda x: x[0])
    return "\n\n".join(r[1] for r in results)


async def _docx_embedded_image_text(path: str) -> str:
    """
    Extract text from images embedded inside a DOCX file (word/media/*).
    Returns concatenated vision-extracted text, or "" if none found.
    """
    try:
        with ZipFile(path) as zf:
            image_names = [
                n for n in zf.namelist()
                if n.startswith("word/media/")
                and n.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".webp"))
            ]

        if not image_names:
            return ""

        print(f"[vision-docx] {len(image_names)} embedded image(s) found")

        async def _extract_one(img_name: str) -> str:
            with ZipFile(path) as zf:
                img_bytes = zf.read(img_name)

            img = Image.open(BytesIO(img_bytes))
            b64 = _encode_pil(img)
            return await _vision_text_from_images([b64], context="medical SOP embedded image")

        parts = await asyncio.gather(*[_extract_one(n) for n in image_names])
        return "\n\n".join(p for p in parts if p.strip())

    except Exception as e:
        print(f"[vision-docx] embedded image extraction failed: {e}")
        return ""


# ── Focused parallel extractors ───────────────────────────────────────────────

async def _extract_meta(text: str) -> dict:
    """Extracts: basic_information + provider_information + workflow_process"""
    prompt = f"""Extract ONLY these sections from the medical SOP below.

Return ONLY this JSON (no extra keys):
{{
  "basic_information": {{"sop_title": "", "category": ""}},
  "provider_information": {{
    "billing_provider_name": "", "billing_provider_npi": "",
    "provider_tax_id": "", "billing_address": "",
    "software": "", "clearinghouse": ""
  }},
  "workflow_process": {{
    "workflow_description": "",
    "eligibility_verification_portals": [],
    "posting_charges_rules": []
  }}
}}

Rules:
- Provider info may appear in headers, footers, letterhead, or signature blocks — search ENTIRE document.
- eligibility_verification_portals: list portal names/URLs.
- posting_charges_rules: list rules about how charges are posted.
- Use "" for missing fields, NOT null.

DOCUMENT:
{text[:12000]}"""
    return await _call_ai(prompt, max_tokens=2000)


async def _extract_billing(text: str) -> dict:
    """Extracts: billing_guidelines only"""
    prompt = f"""Extract ONLY billing guidelines from the medical SOP below.

Billing guidelines = operational/documentation rules ONLY.
Examples: authorization requirements, claim submission rules, timely filing, documentation requirements.

STRICT EXCLUSIONS — do NOT include rules that contain:
- CPT codes (numeric like 99213, J0129)
- ICD-10 codes (letter+number like M17.0)
- Payer-specific rules (those belong in payer_guidelines)

Group rules by their heading/category. Infer category name from surrounding text.

Return ONLY this JSON:
{{
  "billing_guidelines": [
    {{
      "category": "<heading from document>",
      "rules": [{{"description": "<exact original text>"}}]
    }}
  ]
}}

DOCUMENT:
{text[:16000]}"""
    return await _call_ai(prompt, max_tokens=4000)


async def _extract_payers(text: str) -> dict:
    """Extracts: payer_guidelines only"""
    prompt = f"""Extract ONLY payer-specific guidelines from the medical SOP below.

MUST extract:
- ERA setup status (e.g. "Completed", "Form Submitted")
- EDI setup info
- Claim mailing addresses
- Network/credentialing status (INN/OON/NA)
- Timely Filing Limits (TFL)
- Payer ID numbers
- Any rule mentioning a specific payer (Medicare, Medicaid, Aetna, BCBS, UHC, Cigna, etc.)

CRITICAL MERGING RULE:
If the same payer appears in multiple tables/sections, merge ALL their data into ONE object.

STRICT ASSOCIATION:
Only link data (payerId, tfl, etc.) to a payer if EXPLICITLY stated for that payer. No proximity guessing.

Return ONLY this JSON:
{{
  "payer_guidelines": [
    {{
      "payerName": "", "description": "", "payerId": "",
      "eraStatus": "", "ediStatus": "", "tfl": "",
      "networkStatus": "", "mailingAddress": ""
    }}
  ]
}}

DOCUMENT:
{text[:16000]}"""
    return await _call_ai(prompt, max_tokens=6000)


async def _extract_coding(text: str) -> dict:
    """
    Extracts CPT rules from narrative text (inline modifier rules, replacement
    rules, unit caps, admin code rules) AND all ICD rules (replace-with,
    do-not-bill-together, use-only-when).
    Table rows are handled separately by _extract_coding_from_tables.
    """
    prompt = f"""You are extracting CPT and ICD coding rules from a medical SOP.

CPT RULES  ->  coding_rules_cpt
Extract EVERY sentence or rule that mentions a CPT or HCPCS code.

CPT/HCPCS code formats include:
  Numeric: 99213, 73502, 73521, 77080, 96372, 96365, 96366, 96413, 81002
  J-codes: J0129, J0897, J1010, J1200, J2507, J2919, J3489, J7321, j7050, etc.

For EACH rule create ONE object:
  cptCode       -> main CPT/HCPCS code mentioned (e.g. "J0129", "73502")
  description   -> full rule text exactly as written in the document
  modifier      -> any modifier mentioned (e.g. "JZ", "JA", "LT", "RT", "50", "59", "95", "EJ")
  replacementCPT-> if rule says replace X with Y (e.g. "73521")
  units         -> if a specific unit count is mentioned (e.g. "120 max")
  ndcCode       -> leave blank (NDC codes are in tables, not narrative)
  chargePerUnit -> leave blank (charges are in tables, not narrative)

Examples:
  "Please use JZ modifier in all J code CPTs"  -> cptCode:"J-codes", modifier:"JZ"
  "JA modifier should be only use with Medicare for CPT J0129" -> cptCode:"J0129", modifier:"JA"
  "CPT 73502 marked on super bill with LR & RT then use 73521" -> cptCode:"73502", replacementCPT:"73521"
  "CPT J1010 maximum can be billed with 120 units only" -> cptCode:"J1010", units:"120 max"
  "If J7321 is billed more than once, use modifier EJ" -> cptCode:"J7321", modifier:"EJ"
  "For CPT J0897 use admin code 96372 with Aetna" -> TWO objects: J0897 and 96372

ICD RULES  ->  coding_rules_icd
Extract EVERY sentence or rule mentioning an ICD-10 diagnosis code.

ICD-10 format: starts with a LETTER followed by numbers (M17.0, Z00.00, L93.0, M54.50, M1A9XX1)

For EACH ICD code create ONE object:
  icdCode     -> the ICD code exactly as written (e.g. "M17.0", "M54.50")
  description -> full rule text exactly as written
  notes       -> key instruction (e.g. "Replace with M54.59", "Do not bill with M05.89")

CRITICAL: If one sentence mentions multiple ICD codes, create a SEPARATE object for each:
  "M54.50 replace with M54.59"   -> icdCode:"M54.50", notes:"Replace with M54.59"
  "M25.50 replace with M25.59"   -> icdCode:"M25.50", notes:"Replace with M25.59"
  "Do not bill both M45.9 and M05.89 together" -> TWO objects: M45.9 AND M05.89
  "Use M17.0 only when M17.12, M17.9 & M17.11 is given" -> icdCode:"M17.0"

NEVER put CPT codes in coding_rules_icd.
NEVER put ICD codes in coding_rules_cpt.

Return ONLY valid JSON (no markdown, no extra text):
{{
  "coding_rules_cpt": [
    {{"cptCode": "", "description": "", "ndcCode": "", "units": "", "chargePerUnit": "", "modifier": "", "replacementCPT": ""}}
  ],
  "coding_rules_icd": [
    {{"icdCode": "", "description": "", "notes": ""}}
  ]
}}

DOCUMENT:
{text[:20000]}"""
    return await _call_ai(prompt, max_tokens=12000)


async def _extract_coding_from_tables(text: str) -> dict:
    """
    Extracts ALL CPT/HCPCS rows from structured tables.
    Handles BOTH infusion/drug NDC tables AND x-ray modifier/replacement tables.
    """
    table_text = _extract_table_sections(text)
    if not table_text.strip():
        return {"coding_rules_cpt": [], "coding_rules_icd": []}

    prompt = f"""Extract ALL CPT and HCPCS code rows from these medical document tables.

There are TWO types of tables — handle both:

TYPE 1 — Infusion/Drug table (columns: CPT Code | Description | NDC Code | Units | Charge per Unit)
  For each row:
  - cptCode      -> J-code or CPT (e.g. "J0129", "J1010")
  - description  -> drug/procedure name (e.g. "ORENCIA (Injection, abatacept, 10Mg)")
  - ndcCode      -> NDC code exactly as shown (e.g. "00003-2187-13")
  - units        -> unit count or "As per received superbill"
  - chargePerUnit-> charge amount (e.g. "$100.00", "$3750.00")
  - modifier     -> leave blank
  - replacementCPT -> leave blank

TYPE 2 — X-ray/Modifier table (columns: CPT Code | Modifier | Replacement CPT Code)
  For each row:
  - cptCode      -> the CPT code (e.g. "73600", "73510", "72040")
  - description  -> "X-ray code"
  - modifier     -> modifier(s) listed (e.g. "50/LT/RT", "NO Modifier", "Deleted code")
  - replacementCPT -> replacement code if listed (e.g. "73521"), else blank
  - ndcCode, units, chargePerUnit -> leave blank

Rules:
- Extract EVERY row — do NOT skip any
- Preserve exact values (NDC codes, dollar amounts, unit counts)
- Do NOT hallucinate rows not present in the table

Return ONLY valid JSON (no markdown):
{{
  "coding_rules_cpt": [
    {{"cptCode": "", "description": "", "ndcCode": "", "units": "", "chargePerUnit": "", "modifier": "", "replacementCPT": ""}}
  ],
  "coding_rules_icd": []
}}

TABLES:
{table_text}"""
    result = await _call_ai(prompt, max_tokens=8000)
    # Repair J-codes where pdfminer dropped the leading "J" character
    for row in result.get("coding_rules_cpt", []):
        if row.get("cptCode"):
            row["cptCode"] = _repair_j_code(row["cptCode"])
    return result


def _extract_table_sections(text: str) -> str:
    tables = []
    lines = text.splitlines()
    capture = False
    current: list = []
    for line in lines:
        if "TABLE" in line and "START" in line:
            capture = True
            current = []
        elif "TABLE" in line and "END" in line:
            capture = False
            tables.append("\n".join(current))
        elif capture:
            current.append(line)
    return "\n\n".join(tables)


class AISOPService:

    # ── Text extraction ──────────────────────────────────────────────────────

    @staticmethod
    def extract_pdf_text(path: str) -> str:
        return extract_text(path)

    @staticmethod
    def process_sop_extraction(
        sop_id: str,
        file_path: str = None,
        content_type: str = None,
        s3_key: str = None,
    ):
        from app.core.database import SessionLocal
        from app.services.s3_service import s3_service
        import tempfile

        db = SessionLocal()
        temp_file_to_remove = None

        try:
            if s3_key and not file_path:
                file_extension = s3_key.split(".")[-1] if "." in s3_key else ""
                with tempfile.NamedTemporaryFile(
                    delete=False,
                    suffix=f".{file_extension}" if file_extension else "",
                ) as tmp:
                    file_data = asyncio.run(s3_service.download_file(s3_key))
                    tmp.write(file_data)
                    file_path = tmp.name
                    temp_file_to_remove = file_path

            if not file_path:
                raise Exception("No file source provided for extraction")

            text = asyncio.run(AISOPService.extract_text(file_path))
            structured = asyncio.run(AISOPService.ai_extract_sop_structured(text))

            active_status = db.query(Status).filter(
                Status.code == "ACTIVE", Status.type == "GENERAL"
            ).first()

            sop = db.query(SOP).filter(SOP.id == sop_id).first()
            if not sop:
                return

            structured = AISOPService.normalize_ai_sop(structured)

            sop.title = (
                structured.get("basic_information", {}).get("sop_title") or sop.title
            )
            sop.category = (
                structured.get("basic_information", {}).get("category") or sop.category
            )
            provider_info = structured.get("provider_information") or {}

            if not provider_info.get("billingProviderNPI") and sop.client:
                provider_info["billingProviderNPI"] = sop.client.npi

            sop.provider_info = provider_info
            sop.workflow_process = structured.get("workflow_process")

            from app.models.sop import SOPDocument

            doc = db.query(SOPDocument).filter(
                SOPDocument.sop_id == sop_id,
                SOPDocument.s3_key == s3_key if s3_key else None,
            ).first()

            if not doc:
                doc = db.query(SOPDocument).filter(
                    SOPDocument.sop_id == sop_id,
                    SOPDocument.category == "Source file",
                ).first()

            if doc:
                source_name = "source_file" if doc.category == "Source file" else (doc.name or "Document")

                def inject_source(items):
                    if not items:
                        return items
                    for item in items:
                        if isinstance(item, dict):
                            if "rules" in item:
                                for r in item["rules"]:
                                    r["source"] = source_name
                            else:
                                item["source"] = source_name
                    return items

                if doc.category == "Source file":
                    sop.billing_guidelines = structured.get("billing_guidelines")
                    sop.payer_guidelines = structured.get("payer_guidelines")
                    sop.coding_rules_cpt = structured.get("coding_rules_cpt")
                    sop.coding_rules_icd = structured.get("coding_rules_icd")
                else:
                    doc.billing_guidelines = inject_source(structured.get("billing_guidelines"))
                    doc.payer_guidelines = inject_source(structured.get("payer_guidelines"))
                    doc.coding_rules_cpt = inject_source(structured.get("coding_rules_cpt"))
                    doc.coding_rules_icd = inject_source(structured.get("coding_rules_icd"))

                doc.processed = True

            sop.status_id = active_status.id if active_status else sop.status_id
            db.commit()

        except Exception as e:
            print(f"Error in background extraction: {e}")
            failed_status = db.query(Status).filter(
                Status.code == "FAILED", Status.type == "DOCUMENT"
            ).first()
            sop = db.query(SOP).filter(SOP.id == sop_id).first()
            if sop and failed_status:
                sop.status_id = failed_status.id
                db.commit()

        finally:
            db.close()
            target_path = temp_file_to_remove or file_path
            if target_path and os.path.exists(target_path):
                try:
                    os.remove(target_path)
                except Exception:
                    pass

    @staticmethod
    async def extract_image_text(path: str) -> str:
        with Image.open(path) as img:
            buffered = BytesIO()
            img.convert("RGB").save(buffered, format="JPEG", quality=92)
            image_bytes = buffered.getvalue()

        image_base64 = base64.b64encode(image_bytes).decode("utf-8")

        response = await openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": "You are an OCR engine. Extract ALL visible text exactly as-is. No summaries.",
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": "Extract all text. Preserve line breaks.",
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{image_base64}",
                                "detail": "high",
                            },
                        },
                    ],
                },
            ],
            temperature=0,
            max_tokens=4000,
        )

        text = response.choices[0].message.content.strip()
        if not text:
            raise HTTPException(422, "No text extracted from image")
        return text

    @staticmethod
    def extract_docx_text(path: str) -> str:
        doc = Document(path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table_index, table in enumerate(doc.tables, start=1):
            parts.append(f"\n--- TABLE {table_index} START ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
            parts.append(f"--- TABLE {table_index} END ---\n")
        return "\n".join(parts)

    @staticmethod
    def extract_excel_text(path: str) -> str:
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == ".xls":
                df_dict = pd.read_excel(path, sheet_name=None, engine="xlrd")
            elif ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
                df_dict = pd.read_excel(path, sheet_name=None, engine="openpyxl")
            else:
                raise ValueError(f"Unsupported Excel extension: {ext}")
            full_text = ""
            for sheet_name, df in df_dict.items():
                full_text += f"\n--- Sheet: {sheet_name} ---\n"
                full_text += df.fillna("").astype(str).to_string(index=False)
            return full_text
        except Exception as e:
            raise Exception(f"Excel extraction failed: {str(e)}")

    @staticmethod
    async def extract_text(path: str) -> str:
        ext = os.path.splitext(path)[1].lower()

        # ── PDF ──────────────────────────────────────────────────────────────
        if ext == ".pdf":
            text = pdf_extract(path)

            if _is_garbled_text(text):
                # Font-encoded / scanned PDF — pdfminer output is unusable.
                # Fall back to page-image vision extraction via GPT-4o.
                print(
                    f"[extract_text] PDF text garbled (font encoding issue) — "
                    f"switching to vision extraction for {os.path.basename(path)}"
                )
                text = await _pdf_to_vision_text(path)

            return text

        # ── DOCX ─────────────────────────────────────────────────────────────
        if ext == ".docx":
            doc = Document(path)
            parts = []

            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())

            # Tables with TABLE markers so _extract_table_sections() can find them
            for table_index, table in enumerate(doc.tables, start=1):
                parts.append(f"\n--- TABLE {table_index} START ---")
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
                parts.append(f"--- TABLE {table_index} END ---\n")

            text = "\n".join(parts)

            # Embedded images (superbill screenshots, scanned inserts, etc.)
            embedded_text = await _docx_embedded_image_text(path)
            if embedded_text:
                print(
                    f"[extract_text] Appending vision text from "
                    f"{len(embedded_text.split(chr(10)))} lines of embedded images"
                )
                text = text + "\n\n--- EMBEDDED IMAGE CONTENT ---\n" + embedded_text

            return text

        # ── Excel ─────────────────────────────────────────────────────────────
        if ext in [".xlsx", ".xlsm", ".xltx", ".xltm"]:
            wb = load_workbook(path, data_only=True)
            parts = []
            for sheet in wb.worksheets:
                parts.append(f"\n--- SHEET: {sheet.title} ---")
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(cell) for cell in row if cell is not None
                    )
                    if row_text.strip():
                        parts.append(row_text)
            return "\n".join(parts)

        if ext == ".xls":
            df_dict = pd.read_excel(path, sheet_name=None, engine="xlrd")
            parts = []
            for sheet_name, df in df_dict.items():
                parts.append(f"\n--- SHEET: {sheet_name} ---")
                parts.append(df.fillna("").astype(str).to_string(index=False))
            return "\n".join(parts)

        # ── Images ────────────────────────────────────────────────────────────
        if ext in [".png", ".jpg", ".jpeg"]:
            return await AISOPService.extract_image_text(path)

        raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def normalize_ai_sop(data: dict) -> dict:
        category = (
            data.get("basic_information", {}).get("category")
            if isinstance(data.get("basic_information"), dict)
            else data.get("category")
        )
        if isinstance(category, dict):
            category = category.get("title", "")

        if "basic_information" not in data:
            data["basic_information"] = {}

        data["basic_information"]["category"] = category or data.get("category", "")
        data["category"] = data["basic_information"]["category"]

        source = "AI Extracted"

        provider_info_raw = data.get("provider_information") or {}
        data["provider_information"] = {
            "providerName": provider_info_raw.get("billing_provider_name")
            or provider_info_raw.get("providerName"),
            "billingProviderName": provider_info_raw.get("billing_provider_name")
            or provider_info_raw.get("billingProviderName"),
            "billingProviderNPI": provider_info_raw.get("billing_provider_npi")
            or provider_info_raw.get("billingProviderNPI"),
            "providerTaxID": provider_info_raw.get("provider_tax_id")
            or provider_info_raw.get("providerTaxID"),
            "billingAddress": provider_info_raw.get("billing_address")
            or provider_info_raw.get("billingAddress"),
            "software": provider_info_raw.get("software"),
            "clearinghouse": provider_info_raw.get("clearinghouse"),
        }

        # workflow_raw = data.get("workflow_process") or {}
        # data["workflow_process"] = {
        #     "description": workflow_raw.get("workflow_description")
        #     or workflow_raw.get("description"),
        #     "eligibility_verification_portals": workflow_raw.get(
        #         "eligibility_verification_portals", []
        #     ),
        #     "posting_charges_rules": workflow_raw.get("posting_charges_rules", []),
        # }
        workflow_raw = data.get("workflow_process") or {}

        rules = workflow_raw.get("posting_charges_rules", [])

        normalized_rules = []
        for r in rules:
            if isinstance(r, str):
                normalized_rules.append(r)
            elif isinstance(r, dict):
                normalized_rules.append(r.get("description", ""))

        data["workflow_process"]["posting_charges_rules"] = normalized_rules
        guidelines = data.get("billing_guidelines", [])
        normalized_guidelines = []
        for group in guidelines:
            if not isinstance(group, dict):
                continue
            normalized_guidelines.append(
                {
                    "category": group.get("category", "Guidelines"),
                    "rules": [
                        {"description": r.get("description", ""), "source": source}
                        for r in group.get("rules", [])
                        if isinstance(r, dict)
                    ],
                }
            )
        data["billing_guidelines"] = normalized_guidelines

        payer_guidelines = data.get("payer_guidelines", [])
        normalized_payers = []
        for pg in payer_guidelines:
            if isinstance(pg, str):
                normalized_payers.append(
                    {
                        "payerName": "Unknown",
                        "description": pg,
                        "payerId": "",
                        "eraStatus": "",
                        "tfl": "",
                        "networkStatus": "",
                        "mailingAddress": "",
                        "source": source,
                    }
                )
            elif isinstance(pg, dict):
                normalized_payers.append(
                    {
                        "payerName": pg.get("payerName")
                        or pg.get("payer_name")
                        or pg.get("title")
                        or pg.get("payer")
                        or "Unknown",
                        "description": pg.get("description") or "",
                        "payerId": pg.get("payerId") or pg.get("payer_id") or "",
                        "eraStatus": pg.get("eraStatus") or pg.get("era_status") or "",
                        "ediStatus": pg.get("ediStatus") or pg.get("edi_status") or "",
                        "tfl": pg.get("tfl") or "",
                        "networkStatus": pg.get("networkStatus")
                        or pg.get("network_status")
                        or "",
                        "mailingAddress": pg.get("mailingAddress")
                        or pg.get("mailing_address")
                        or pg.get("address")
                        or "",
                        "source": pg.get("source") or source,
                    }
                )
        data["payer_guidelines"] = normalized_payers

        for key in ("coding_rules_cpt", "coding_rules_icd"):
            data[key] = [
                {**r, "source": source}
                for r in data.get(key, [])
                if isinstance(r, dict)
            ]

        return data

    @staticmethod
    def safe_parse_json(raw: str) -> dict:
        raw = raw.strip()
        if raw.startswith("```"):
            raw = raw.replace("```json", "").replace("```", "").strip()
        if not raw.startswith("{"):
            raise ValueError("AI did not return JSON")
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            last_brace = raw.rfind("}")
            if last_brace != -1:
                try:
                    return json.loads(raw[: last_brace + 1])
                except ValueError as e:
                    raise HTTPException(422, str(e))
                except Exception:
                    raise HTTPException(500, "Internal server error")
            raise ValueError(f"Invalid JSON from AI:\n{raw[:1000]}")

    @staticmethod
    def extract_table_sections(text: str) -> str:
        return _extract_table_sections(text)

    # ── Legacy compat ────────────────────────────────────────────────────────

    @staticmethod
    async def _call_ai(prompt: str) -> dict:
        return await _call_ai(prompt)

    @staticmethod
    async def extract_narrative_and_rules(text: str) -> dict:
        meta, billing, payers, coding = await asyncio.gather(
            _extract_meta(text),
            _extract_billing(text),
            _extract_payers(text),
            _extract_coding(text),
        )
        return {**meta, **billing, **payers, **coding}

    @staticmethod
    async def extract_tables(text: str) -> dict:
        return await _extract_coding_from_tables(text)

    # ── Main entry point ─────────────────────────────────────────────────────

    @staticmethod
    async def ai_extract_sop_structured(text: str) -> dict:
        """
        Fires all 5 extraction calls simultaneously via asyncio.gather.
        Total time ≈ slowest single call (~5-10s) instead of sum of all calls.

        Calls:
          1. _extract_meta        → title, provider, workflow
          2. _extract_billing     → billing guidelines
          3. _extract_payers      → payer guidelines
          4. _extract_coding      → CPT + ICD from narrative text
          5. _extract_coding_from_tables → CPT from table sections
        """
        (
            meta,
            billing,
            payers,
            coding_narrative,
            coding_tables,
        ) = await asyncio.gather(
            _extract_meta(text),
            _extract_billing(text),
            _extract_payers(text),
            _extract_coding(text),
            _extract_coding_from_tables(text),
        )

        # Merge + deduplicate CPT from narrative and table passes
        cpt_combined: list = (
            coding_narrative.get("coding_rules_cpt", [])
            + coding_tables.get("coding_rules_cpt", [])
        )
        seen_cpt: set = set()
        deduped_cpt = []
        for r in cpt_combined:
            key = r.get("cptCode") or r.get("description", "")
            if key and key not in seen_cpt:
                seen_cpt.add(key)
                deduped_cpt.append(r)
            elif not key:
                deduped_cpt.append(r)

        final = {
            **meta,
            **billing,
            **payers,
            "coding_rules_cpt": deduped_cpt,
            "coding_rules_icd": coding_narrative.get("coding_rules_icd", []),
        }

        return AISOPService.normalize_ai_sop(final)


# ── Medical Document Type Detection & Extraction ──────────────────────────────
#
# This section handles per-page AI extraction for scanned medical documents
# that arrive as multi-page PDFs.  Each page (or page-pair) is a distinct
# document type that needs its own extraction logic:
#
#   SUPERBILL       – checked CPT codes, patient header, return instructions
#   ICD10_SHEET     – checked diagnosis codes + ICD-10 for the visit
#   DXA_SHEET       – DXA order with patient, CPT, and diagnosis selection
#   INSURANCE_CARD  – payer name, member ID, group, claims address
#   FINANCIAL_FORM  – financial responsibility agreement (signature / date)
#   UNKNOWN         – anything else
#
# Entry point:  extract_medical_document_pages(path) → list[dict]
# Each dict is a fully-structured page record.  The caller can persist these
# directly or stitch them together into a visit record.


async def _call_vision_json(
    images_b64: list[str],
    system_prompt: str,
    user_prompt: str,
    max_tokens: int = 4000,
) -> dict:
    """
    Send 1-2 page images to GPT-4o with a focused JSON extraction prompt.
    Returns parsed dict (3-tier fallback same as _call_ai).
    """
    content: list[dict] = [{"type": "text", "text": user_prompt}]
    for b64 in images_b64:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/jpeg;base64,{b64}", "detail": "high"},
        })

    response = await openai_client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content},
        ],
        temperature=0,
        max_tokens=max_tokens,
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()
    start = raw.find("{")
    end = raw.rfind("}")
    if start == -1:
        raise ValueError(f"Vision returned non-JSON: {raw[:200]}")
    raw = raw[start:] if (end == -1 or end < start) else raw[start: end + 1]

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass
    try:
        from json_repair import repair_json
        return json.loads(repair_json(raw))
    except Exception:
        pass
    last = raw.rfind("}")
    if last > 0:
        try:
            return json.loads(raw[:last + 1])
        except Exception:
            pass
    raise ValueError(f"Cannot parse vision JSON: {raw[:300]}")


# ── Specialist extractor registry ─────────────────────────────────────────────


def _build_classifier_prompt(schemas: list[dict]) -> tuple[str, str]:
    """
    Build a dynamic classifier system prompt + user prompt from DB schemas.

    Each schema dict must have at minimum:
      - type_name:   str   (e.g. "SUPERBILL")
      - description: str   (human-readable description of what this doc looks like)
      - fields:      list  (extraction_fields from the template)

    Returns (system_prompt, user_prompt).
    """
    type_names = [s["type_name"].strip().upper() for s in schemas]
    valid_types = " | ".join(type_names) + " | UNKNOWN"

    # Build per-type description block
    type_blocks = []
    for s in schemas:
        name = s["type_name"].strip().upper()
        desc = (s.get("description") or "").strip()
        fields = s.get("fields") or []

        # Pull field names as visual hints (what the model should look for on-page)
        field_hints = []
        for f in fields[:12]:   # cap at 12 hints to keep prompt lean
            fn = f.get("fieldName") or f.get("name") or ""
            if fn:
                field_hints.append(fn)

        hint_str = ""
        if field_hints:
            hint_str = "\n  Key fields visible on this document type: " + ", ".join(field_hints)

        block = f"{name}\n  {desc}{hint_str}" if desc else f"{name}{hint_str}"
        type_blocks.append(block)

    types_section = "\n\n".join(type_blocks)

    system = (
        "You are a medical document page classifier. "
        "Respond with valid JSON only. No markdown, no explanation."
    )

    user = f"""Classify this medical document page image.

Return ONLY this JSON:
{{
  "doc_type": "<one of: {valid_types}>",
  "confidence": "<HIGH | MEDIUM | LOW>",
  "notes": "<one short sentence explaining the key visual signal you used>"
}}

CLASSIFICATION RULES:
- Match the page to the single best doc_type from the list below.
- Use the description and key fields as your visual guide.
- If the page clearly does not match any listed type, use UNKNOWN.
- Confidence HIGH = multiple strong signals present; MEDIUM = partial match; LOW = guess.

DOCUMENT TYPES:

{types_section}

UNKNOWN
  Use when the page does not match any of the listed types.
  Examples: blank pages, fax cover sheets, referral letters, standalone lab reports."""

    return system, user


async def _classify_page_dynamic(b64: str, schemas: list[dict]) -> str:
    """Classify one page image against the provided DB schemas."""
    system, user = _build_classifier_prompt(schemas)
    try:
        result = await _call_vision_json([b64], system, user, max_tokens=200)
        raw_type = result.get("doc_type", "UNKNOWN").strip().upper()
        # Validate: must be one of the schema type names or UNKNOWN
        valid = {s["type_name"].strip().upper() for s in schemas} | {"UNKNOWN"}
        return raw_type if raw_type in valid else "UNKNOWN"
    except Exception as e:
        print(f"[classify_page_dynamic] error: {e}")
        return "UNKNOWN"



# ── Per-type specialist extractors ────────────────────────────────────────────


async def _extract_superbill(b64: str) -> dict:
    """
    Extract structured data from a Greater Washington Arthritis superbill page.

    This superbill has a fixed printed template with six sections:
      A. OFFICE VISITS    — CPT 99201-99215
      B. CONSULTATION     — CPT 99241-99245
      C. PROCEDURES       — CPT 20600, 20605, 20610, 64405, 20526, 20550, 71045, 20552, 20553
      D. MEDICATIONS      — J-codes: J0702, J1030/J1040, J1100, J3301, J3304, + many more
      E. RADIOLOGY        — CPT 73600-73660, 72xxx, 77080, 20604-20611, 76942, 76882, 73020-73030 etc.
      F. L-CODES          — L3931, L3933, L3924, L3908, L3510, L30S0, LI902, L1851, S8451 (×2), A4467, A6530
      G. LAB FEE          — 36415

    Returns structured data including all MARKED items only.
    """
    prompt = """Extract ALL data from this medical superbill image.

PATIENT HEADER (top of page — always handwritten or typed):
- patient_name: full name written at top
- dob: date of birth (MM/DD/YYYY)
- date_of_service: visit date (MM/DD/YYYY)
- insurance: insurance company name(s) written at top right
- provider: name after "Provider" label (e.g. "DR MONSHIE", "HASHEM MONSHIE, MD", "MARTIN PA-C")
- balance: dollar amount in Balance field (may be blank)
- copay: dollar amount in Copay field (may be blank)
- today_payment: dollar amount in Today's Payment field (may be handwritten)
- return_visit: {value: "N", unit: "Days|Weeks|Months"} — read the RETURN line at bottom

MARKED ITEMS — only include items that are VISUALLY SELECTED:
A selected item means ANY of: checkbox is checked ☑, item is circled, item is underlined,
handwriting appears directly next to/on it, or a mark/scribble is present adjacent to it.
Do NOT include items that have only a printed checkbox □ with no mark.

SUPERBILL SECTIONS to scan:
- A. OFFICE VISITS: New Patient (99201-99205) and Established Patient (99211-99215)
- B. CONSULTATION: Outpatient 1-5 (99241-99245)
- C. PROCEDURES AND SERVICES: arthrocentesis codes, CXR, myofascial inj, nerve block, carpal tunnel, trigger finger (20600, 20605, 20610, 71045, 20552, 20553, 64405, 20526, 20550)
- D. MEDICATIONS: Celestone, Depomedrol, Dexamethasone, Kenalog, Zilretta, Actemra, B12, Benedryl, Benlysta, Boniva, Cimzia, Cosentyx, Euflexxa, Enbrel, Evenity, infusion codes, flu shot, Humira, Hyalgan, Ilaris, Inflectra, Kevzara, Krystexxa, Methotrexate, Orencia, Orthovisc, PPD, Prolia, Reclast, Remicade, Renflexis, Rituxan, RPR, Salin, Saphnelo, Simponi, Skyrizi, Solumedrol, Stelara, Subcutaneous Biologics, Supartz, Synvisc, Taltz, Toradol, Tremfya, Truxima, Zoledronic, Zofran (J-codes and procedure codes)
- E. RADIOLOGY: ankle, cervical spine, collar bone, CXR, DXA scan, elbow, fingers, femur, foot, forearm, hand, heel, hip, humerus, knees, L-spine, lower leg, pelvic, sacroiliac, sacrum coccyx, sono arthrocentesis (small/medium/large joint, needle, evaluation with image), shoulder, thoracic spine, toes, wrist, lumbar, hips & pelvis (73xxx, 72xxx, 20604-20611, 76942, 76882, 73020-73030) — for bilateral items capture which side(s): L and/or R
- F. L-CODES: arthritis glove, finger orthosis, thumb stabilizer, carpal tunnel brace, plantar fasciitis, foot arch support, ankle/foot compression brace, knee compression brace, wrist splint, ankle splint, tennis elbow brace, gradient compression stocking
- G. LAB FEE: 36415

For RADIOLOGY items that have L/R laterality: capture as separate entries if both sides marked, or include side in notes.
For MEDICATIONS with handwritten dosage annotations: include the handwritten dose/quantity in notes.
For items where the CPT code itself is circled or underlined, treat as selected.

Return ONLY this JSON (no markdown):
{
  "patient_name": "",
  "dob": "",
  "date_of_service": "",
  "insurance": "",
  "provider": "",
  "balance": "",
  "copay": "",
  "today_payment": "",
  "return_visit": {"value": "", "unit": ""},
  "selected_cpts": [
    {
      "code": "",
      "description": "",
      "section": "OFFICE_VISIT|CONSULTATION|PROCEDURE|MEDICATION|RADIOLOGY|L_CODE|LAB",
      "modifier": "",
      "notes": ""
    }
  ]
}"""

    return await _call_vision_json(
        [b64],
        "You are a precise medical billing data extractor for Greater Washington Arthritis, "
        "Rheumatology and Osteoporosis Center. Return valid JSON only. "
        "Never include unchecked/unmarked items.",
        prompt,
        max_tokens=4000,
    )


async def _extract_icd10_sheet(b64: str) -> dict:
    """
    Extract checked/marked diagnosis codes from the ICD-10 diagnosis sheet.

    This is a 3-column alphabetical diagnosis list (ADHD → Xerostomia) with CPT codes
    in a header column. The right side has large-print administrative labels:
    Photo / Referral/PCP / Email / Review / Insurance Data Entry / Signed out By.

    Marked items can be:
    - Checkbox checked ☑
    - Item circled
    - Item underlined
    - Handwritten text or code correction next to it
    - Strikethrough with amended code written nearby
    """
    prompt = """Examine this ICD-10 diagnosis sheet image carefully.

STEP 1 — Scan ALL three columns from top to bottom looking for ANY visual mark.
STEP 2 — For each marked item, record it in the output.

A mark includes: checked checkbox ☑, circle around text or code, underline,
strikethrough, pen/pencil scribble, or handwritten correction.

AMENDMENT RULE: If an item is crossed out with a new code written nearby, record:
  status = "amended", original icd_code = the printed code, amended_code = the handwritten code.

SELECTED RULE: If checked or circled with no correction: status = "selected".

CROSSED_OUT RULE: If struck through with no replacement: status = "crossed_out".

Also extract:
- signed_out_by: the name written on "Signed out By ___" line (if legible)
- referral_pcp: text written on "Referral / PCP:" line (if present)
- review_notes: text written on "Review:" line (if present)

Return ONLY this JSON (no markdown):
{
  "selected_diagnoses": [
    {
      "icd_code": "",
      "description": "",
      "status": "selected|crossed_out|amended",
      "amended_code": ""
    }
  ],
  "signed_out_by": "",
  "referral_pcp": "",
  "review_notes": ""
}

IMPORTANT: Be thorough — scan every row in all three columns. A faint checkmark or
small circle still counts. If nothing is marked anywhere, return an empty list."""

    return await _call_vision_json(
        [b64],
        "You are a precise medical coding data extractor. "
        "Scan every row carefully for visual marks. Return valid JSON only.",
        prompt,
        max_tokens=4000,
    )


async def _extract_dxa_sheet(b64: str) -> dict:
    """
    Extract data from a DXA Diagnostic Codes order form.

    DXA sheets come in PAIRS:
    - Page 1 (front): practice letterhead, patient name/account#/date/insurance,
      scan type checkboxes (77080/77081/77082/DXA-TBCA), and first half of diagnosis
      categories (Adrenal, Deficiency, Density, Endocrine, Fractures, Gastrointestinal,
      Miscellaneous)
    - Page 2 (back): continuation of diagnosis categories (Ovary, Osteoporosis/Osteopenia,
      Parathyroid, Pituitary, Renal, Testis, Thyroid, E Codes/Drugs, V Codes,
      Codes for DXA-TBCA) plus Follow-up line

    Both pages use this extractor. Fields not present on a page should be null/empty.

    SCAN CPT CODE MEANINGS:
    - Axial Skeleton = 77080
    - Forearm = 77081
    - VFA = 77082
    - DXA Total Body Composition Analysis (DXA-TBCA) = 77080 DXA-TBCA
    """
    prompt = """Extract data from this DXA Diagnostic Codes form page.

This form may be PAGE 1 (has patient info + letterhead + scan type checkboxes) or
PAGE 2 (continuation of diagnosis codes only, no patient info).

For PAGE 1, extract:
- patient_name (from "Patient's Name" line)
- account_number (from "Account #" line, may be blank)
- date (from "Date:" field, format MM/DD/YYYY)
- referring_physician (from "Referring Physician" line, may be blank)
- insurance (from "Patient's Insurance" line, may be blank)
- scan_cpt: which scan checkbox is checked — one of:
    "77080" (Axial Skeleton 77080)
    "77081" (Forearm 77081)
    "77082" (VFA 77082)
    "77080-TBCA" (DXA Total Body Composition Analysis)
    "" (none checked)
- handwritten_number: any large handwritten number in top-left corner (e.g. "184/65", "175/64") — this appears to be height/weight or a reference number

For BOTH pages, extract ALL checked diagnosis codes:
- selected_diagnoses: list of every checked/marked item

The diagnosis categories present on these pages include:
PAGE 1: Adrenal (255.0, 255.3, V56.65), Deficiency (268.2, 268.9, etc.), Density (733.41),
  Endocrine (259.9, 259.3), Fractures (805.0-821.3), Gastrointestinal (271.3-558.9),
  Miscellaneous (203.0, 259.3, 259.9, 307.1)
PAGE 2: Ovary (256.2-758.6), Osteoporosis/Osteopenia (733.00-781.91),
  Parathyroid (252.0-588.81), Pituitary (227.3-253.3), Renal (585.0, 588.0),
  Testis (257.1, 257.2), Thyroid (242.00-242.80),
  More bone/rheumatic (337.20, 710.0, 714.0, 731.0, 731.8, 756.50),
  E Codes/Drugs (E932.0, E932.7, E936.1, E944.4),
  V Codes (V07.4, V49.81, V58.61, V58.65, V58.69, V67.51, V67.59, V82.81, V82.89),
  Codes for DXA-TBCA (278.00, 278.01, 277.7, 783.1)
- Follow-up: text from "Follow up ____" line (if present)

Return ONLY this JSON (no markdown):
{
  "patient_name": "",
  "account_number": "",
  "date": "",
  "referring_physician": "",
  "insurance": "",
  "scan_cpt": "",
  "handwritten_number": "",
  "selected_diagnoses": [
    {
      "code": "",
      "description": "",
      "category": ""
    }
  ],
  "follow_up": ""
}

RULES:
- Only include diagnoses that are CHECKED (box checked ☑ or circled/marked)
- Capture the printed ICD code exactly as shown (old-style ICD-9 format like 733.00 or new like M81.0)
- category = the section heading (e.g. "Osteoporosis", "Fractures", "Thyroid")
- If this is page 2 (no patient header), leave patient fields as ""
- The handwritten number top-left corner of page 1 should always be captured"""

    return await _call_vision_json(
        [b64],
        "You are a precise medical data extractor specializing in DXA order forms. "
        "Return valid JSON only.",
        prompt,
        max_tokens=3000,
    )


async def _extract_insurance_card(b64: str) -> dict:
    """
    Extract insurance card(s), ID document, and financial responsibility form data.

    This page can contain any combination of:
    - Medicare card (red/white/blue, shows Medicare number, Hospital Part A / Medical Part B)
    - Commercial insurance card (Anthem, BCBS, UHC, CareFirst, HealthKeepers, AARP supplement, etc.)
    - Rx/pharmacy benefit card
    - Driver's license or state ID
    - Financial Responsibility Agreement to Pay form (ALWAYS appears on these pages)

    Multiple cards may be shown on one page. Extract ALL of them.
    """
    prompt = """Extract ALL insurance/ID card data and financial responsibility form data from this image.

This page typically shows physical card images (Medicare, commercial insurance, driver's license)
AND a "Financial Responsibility Agreement to Pay" form at the bottom.

INSURANCE CARDS — for each card visible, extract:
Primary card (usually the largest or most prominent):
- payer_name: insurance company name (e.g. "Medicare", "Anthem HealthKeepers Plus",
    "BCBS", "UnitedHealthcare", "CareFirst BlueChoice", "AARP Medicare Supplement",
    "First Health", "Humana", "Cigna")
- plan_type: plan or product name (e.g. "Medicare Advantage", "Supplement Plan F",
    "HMO", "PPO", "Medicaid", "Part A", "Part B")
- member_name: name on card
- member_id: member/subscriber ID number
- group_number: group number if shown
- rx_bin: BIN number (pharmacy benefit)
- rx_pcn: PCN number
- rx_group: Rx group number
- copays: any copay amounts shown (office visit, specialist, ER, etc.)
- claims_address: mailing address for claims
- phone: member services phone number
- effective_date: coverage start date if shown

Secondary card (if a second insurance card is present):
- secondary_payer_name, secondary_member_id, secondary_group_number

ID DOCUMENT (driver's license / state ID if present):
- id_type: "DRIVER_LICENSE" | "STATE_ID"
- id_state: state abbreviation
- id_name: full name on ID
- id_dob: date of birth on ID
- id_number: license/ID number (last 4 digits only for privacy, or full if clearly visible)
- id_expiry: expiration date

FINANCIAL RESPONSIBILITY FORM:
- financial_responsibility_signed: true if a handwritten signature is present on the form
- signature_date: date written next to signature (MM/DD/YYYY)

Return ONLY this JSON (no markdown):
{
  "payer_name": "",
  "plan_type": "",
  "member_name": "",
  "member_id": "",
  "group_number": "",
  "rx_bin": "",
  "rx_pcn": "",
  "rx_group": "",
  "copays": "",
  "claims_address": "",
  "phone": "",
  "effective_date": "",
  "secondary_payer_name": "",
  "secondary_member_id": "",
  "secondary_group_number": "",
  "id_type": "",
  "id_state": "",
  "id_name": "",
  "id_dob": "",
  "id_expiry": "",
  "financial_responsibility_signed": false,
  "signature_date": ""
}"""

    return await _call_vision_json(
        [b64],
        "You are a medical insurance data extractor. "
        "Extract all card data carefully. Return valid JSON only.",
        prompt,
        max_tokens=2000,
    )

async def _extract_financial_form(b64: str) -> dict:
    """
    Extract data from a standalone Financial Responsibility Agreement to Pay form
    (no insurance card present — just the agreement text + signature).
    """
    prompt = """Extract data from this Financial Responsibility Agreement to Pay form.

The form says:
"I acknowledge that the above is my current insurance, and any changes have been notified..."
"I accept FULL FINANCIAL responsibility from Greater Washington Arthritis..."

Return ONLY this JSON (no markdown):
{
  "financial_responsibility_signed": false,
  "signature_date": "",
  "notes": ""
}

- financial_responsibility_signed: true if a handwritten signature line is signed
- signature_date: date written near signature (MM/DD/YYYY format)
- notes: any other handwritten text on the form"""

    return await _call_vision_json(
        [b64],
        "You are a medical document data extractor. Return valid JSON only.",
        prompt,
        max_tokens=500,
    )


# ── Register specialist extractors ────────────────────────────────────────────
# These override the universal extractor for known high-value document types.
# Keys are normalised (uppercase) type names matching DocumentType.name in the DB.
_SPECIALIST_EXTRACTORS = {
    "SUPERBILL":      _extract_superbill,
    "ICD10_SHEET":    _extract_icd10_sheet,
    "DXA_SHEET":      _extract_dxa_sheet,
    "INSURANCE_CARD": _extract_insurance_card,
    "FINANCIAL_FORM": _extract_financial_form,
}


async def _extract_universal(b64: str, schema: dict) -> dict:
    """
    Universal schema-driven extractor.
    Used for any document type that does NOT have a specialist extractor registered.

    Builds the extraction prompt entirely from the DB template fields so that
    new document types work out-of-the-box without any code changes.

    Schema shape expected:
      {
        "type_name":   "SOME_DOC_TYPE",
        "description": "What this document is",
        "fields": [
          {"fieldName": "patientName", "type": "text", "description": "Patient full name"},
          ...
        ]
      }
    """
    type_name = schema.get("type_name", "UNKNOWN").strip().upper()
    description = (schema.get("description") or "").strip()
    fields = schema.get("fields") or []

    if not fields:
        return {"error": f"No fields defined in template for {type_name}"}

    # Build explicit field list with type + description hints
    field_lines = []
    for f in fields:
        fn = f.get("fieldName") or f.get("name") or ""
        ft = f.get("type") or f.get("fieldType") or "text"
        fd = f.get("description") or f.get("label") or ""
        example = f.get("exampleValue") or ""
        hint_parts = [f"type={ft}"]
        if fd:
            hint_parts.append(fd)
        if example:
            hint_parts.append(f'e.g. "{example}"')
        field_lines.append(f'  "{fn}": null   // {", ".join(hint_parts)}')

    fields_block = "\n".join(field_lines)
    field_names_json = json.dumps({(f.get("fieldName") or f.get("name", "")): None for f in fields}, indent=2)

    prompt = f"""Extract structured data from this medical document image.

Document type: {type_name}
{("Description: " + description) if description else ""}

FIELDS TO EXTRACT (extract ALL fields; use null if not found on this page):
{fields_block}

RULES:
- Extract ONLY the fields listed above.
- Preserve exact values as they appear (dates, codes, amounts, names).
- If a field is not visible on this page, set it to null.
- For list/array fields (e.g. line items, diagnoses), return a JSON array.
- Do NOT invent values. Do NOT add extra keys.

Return ONLY this JSON (no markdown, no explanation):
{field_names_json}"""

    return await _call_vision_json(
        [b64],
        f"You are a precise medical data extractor for {type_name} documents. "
        "Return valid JSON only.",
        prompt,
        max_tokens=4000,
    )


# ── Main entry point ──────────────────────────────────────────────────────────

async def extract_medical_document_pages(
    path: str,
    schemas: list[dict] | None = None,
) -> list[dict]:
    """
    Process a multi-page medical document PDF.

    Parameters
    ----------
    path : str
        Path to the PDF file on disk.
    schemas : list[dict] | None
        Document-type schemas fetched from the DB, each with:
          - type_name:   str    normalised doc type name (e.g. "SUPERBILL")
          - description: str    human-readable description used for classification
          - fields:      list   extraction_fields from the active template
        When None or empty, classification falls back to the static specialist
        types only (SUPERBILL, ICD10_SHEET, DXA_SHEET, INSURANCE_CARD, FINANCIAL_FORM).

    Returns
    -------
    list[dict]  — one record per page (or merged DXA pair), each containing:
        page_number, doc_type, encounter_id, data (type-specific dict)

    Usage
    -----
        # With DB schemas (recommended):
        records = await extract_medical_document_pages(path, schemas=schemas)

        # Without schemas (legacy fallback):
        records = await extract_medical_document_pages(path)

        superbills = [r for r in records if r["doc_type"] == "SUPERBILL"]
    """
    from pdf2image import convert_from_path

    # ── Normalise schemas ──────────────────────────────────────────────────
    # Always work with a list; normalise type_name to uppercase
    if schemas:
        normalised_schemas = [
            {**s, "type_name": s["type_name"].strip().upper()}
            for s in schemas
            if s.get("type_name")
        ]
    else:
        # No schemas supplied → build a minimal stub list from the specialist
        # extractor registry so the classifier still has something to work with
        normalised_schemas = [
            {"type_name": k, "description": "", "fields": []}
            for k in _SPECIALIST_EXTRACTORS
        ]

    # schema_map: type_name → schema dict (for extraction phase)
    schema_map: dict[str, dict] = {s["type_name"]: s for s in normalised_schemas}

    loop = asyncio.get_event_loop()

    # Convert pages at 150dpi for better handwriting legibility
    pages: list[Image.Image] = await loop.run_in_executor(
        None,
        lambda: convert_from_path(path, dpi=150),
    )

    if not pages:
        raise ValueError(f"No pages found in {path}")

    print(f"[medical-doc] {len(pages)} pages | {len(normalised_schemas)} schema types: "
          f"{[s['type_name'] for s in normalised_schemas]}")

    # Encode all pages (CPU-bound)
    encoded: list[str] = await loop.run_in_executor(
        None,
        lambda: [_encode_pil(p, quality=85) for p in pages],
    )

    # ── Classify all pages in parallel ────────────────────────────────────
    sem = asyncio.Semaphore(10)

    async def _classify_with_sem(b64: str, page_no: int) -> tuple[int, str]:
        async with sem:
            doc_type = await _classify_page_dynamic(b64, normalised_schemas)
            print(f"[medical-doc] page {page_no}: {doc_type}")
            return page_no, doc_type

    classify_tasks = [_classify_with_sem(b64, i + 1) for i, b64 in enumerate(encoded)]
    classifications: list[tuple[int, str]] = await asyncio.gather(*classify_tasks)

    page_types: dict[int, str] = {pno: dtype for pno, dtype in classifications}

    # ── Extract each page ─────────────────────────────────────────────────
    # Priority: specialist extractor → universal schema extractor → no-op
    async def _extract_page(page_no: int, b64: str) -> dict:
        doc_type = page_types[page_no]
        specialist = _SPECIALIST_EXTRACTORS.get(doc_type)
        schema = schema_map.get(doc_type)

        async with sem:
            try:
                if specialist:
                    # Hand-tuned extractor takes precedence
                    data = await specialist(b64)
                elif schema and schema.get("fields"):
                    # Universal extractor driven by DB template fields
                    data = await _extract_universal(b64, schema)
                else:
                    data = {
                        "note": "No extractor or template fields available for this type",
                        "raw_type": doc_type,
                    }
            except Exception as e:
                print(f"[medical-doc] page {page_no} ({doc_type}) extraction error: {e}")
                data = {"error": str(e)}

        return {
            "page_number": page_no,
            "doc_type": doc_type,
            "data": data,
        }

    extract_tasks = [_extract_page(i + 1, b64) for i, b64 in enumerate(encoded)]
    results: list[dict] = await asyncio.gather(*extract_tasks)

    # Sort by page number
    results.sort(key=lambda r: r["page_number"])

    # ── Post-process: merge consecutive DXA sheet pairs ──────────────────────
    results = _merge_dxa_pairs(results)

    # ── Post-process: group pages by patient encounter ───────────────────────
    results = _tag_patient_encounters(results)

    return results


def _merge_dxa_pairs(results: list[dict]) -> list[dict]:
    """
    Merge consecutive DXA_SHEET page pairs into a single combined record.
    Page 1 contributes patient info + scan_cpt; page 2 adds more diagnoses.
    The merged record keeps page 1's page_number and marks page_2_number.
    """
    merged = []
    i = 0
    while i < len(results):
        rec = results[i]
        if (
            rec["doc_type"] == "DXA_SHEET"
            and i + 1 < len(results)
            and results[i + 1]["doc_type"] == "DXA_SHEET"
        ):
            p1 = rec["data"]
            p2 = results[i + 1]["data"]
            # Merge diagnosis lists
            combined_dx = (p1.get("selected_diagnoses") or []) + (
                p2.get("selected_diagnoses") or []
            )
            merged_data = {**p1}
            merged_data["selected_diagnoses"] = combined_dx
            # Carry over follow_up from page 2 if present
            if p2.get("follow_up"):
                merged_data["follow_up"] = p2["follow_up"]
            merged.append({
                "page_number": rec["page_number"],
                "page_2_number": results[i + 1]["page_number"],
                "doc_type": "DXA_SHEET",
                "data": merged_data,
            })
            i += 2
        else:
            merged.append(rec)
            i += 1
    return merged


def _tag_patient_encounters(results: list[dict]) -> list[dict]:
    """
    Tag each record with an encounter_id.
    A new encounter starts whenever a SUPERBILL page is seen.
    The ICD10_SHEET immediately following a SUPERBILL gets the same encounter_id.
    DXA_SHEET and INSURANCE_CARD records get their own encounter_id if they
    are not adjacent to a SUPERBILL.
    """
    encounter_id = 0
    current_encounter = None
    for rec in results:
        dt = rec["doc_type"]
        if dt == "SUPERBILL":
            encounter_id += 1
            current_encounter = encounter_id
        elif dt == "ICD10_SHEET" and current_encounter is not None:
            pass  # same encounter as the preceding superbill
        else:
            # DXA, insurance cards etc. get their own sequential id
            encounter_id += 1
            current_encounter = encounter_id
        rec["encounter_id"] = current_encounter
    return results


def extract_medical_document_pages_sync(
    path: str,
    schemas: list[dict] | None = None,
) -> list[dict]:
    """
    Synchronous wrapper for use in background threads (e.g. process_sop_extraction).
    Pass schemas from the DB for dynamic classification and extraction.
    """
    return asyncio.run(extract_medical_document_pages(path, schemas=schemas))